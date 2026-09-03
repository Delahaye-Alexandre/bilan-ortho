"""Texte riche des rubriques (lot A du plan « mise en forme ») : analyse et
sérialisation du Markdown restreint, rendu dans les quatre exports, mise en
forme conservée à l'import Word, consigne au modèle et vérificateurs."""
import io
import re
import zipfile

from docx import Document

from app import anonymisation, bilan, export, importer, llm, prompts
from app import texte_riche as tr

# Mêmes échantillons que tests/ui/test_texte_riche_ui.mjs : l'éditeur (JS) et
# les exports (Python) doivent produire la même forme canonique.
ECHANTILLONS = [
    ("Texte brut.", "Texte brut."),
    ("Le test **Alouette** est *chuté* : <u>Compréhension</u> :\n- a\n* b\n• c\nSuite.",
     "Le test **Alouette** est *chuté* : <u>Compréhension</u> :\n\n- a\n- b\n- c\n\nSuite."),
    ("5 * 3 = 15 et (*) note, **non fermé", "5 * 3 = 15 et (*) note, **non fermé"),
    ("Axes :\n1. phono\n2. lecture\n  suite de lecture\n\nParagraphe final.",
     "Axes :\n\n1. phono\n2. lecture\n  suite de lecture\n\nParagraphe final."),
    ("***gras ital*** et **gras *ital* fin**", "***gras ital*** et **gras** ***ital*** **fin**"),
    ("Écart :\n\\- 2 ET à l'Alouette", "Écart :\n\\- 2 ET à l'Alouette"),
    ("", ""),
]


def test_canonique_partagee_avec_l_editeur():
    for entree, attendu in ECHANTILLONS:
        assert tr.canonique(entree) == attendu
        assert tr.canonique(attendu) == attendu  # stable


def test_analyse_segments_et_listes():
    blocs = tr.analyser("Le **test** est <u>chuté</u>.\n- a\n- b\n  suite\n\n1. un\n2. deux")
    assert isinstance(blocs[0], tr.Paragraphe)
    assert [(s.texte, s.gras, s.souligne) for s in blocs[0].segments] == [
        ("Le ", False, False), ("test", True, False), (" est ", False, False),
        ("chuté", False, True), (".", False, False),
    ]
    assert isinstance(blocs[1], tr.Liste) and not blocs[1].ordonnee
    assert [s.texte for s in blocs[1].items[1]] == ["b\nsuite"]
    assert isinstance(blocs[2], tr.Liste) and blocs[2].ordonnee
    assert len(blocs[2].items) == 2


def test_marqueurs_tolerants_et_espaces_hors_marqueurs():
    # Un marqueur mal placé ou non fermé reste du texte.
    assert tr.analyser("a ** b ** c")[0].segments[0].simple
    assert tr.en_clair("**non fermé") == "**non fermé"
    # Un segment gras qui porte une espace de bord la rejette hors des marqueurs.
    assert tr.serialiser_segments([tr.Segment("Alouette ", gras=True)]) == "**Alouette** "
    # Un texte de rubrique existant, brut, est déjà valide et inchangé.
    brut = "Né à terme.\nAudition normale ; pas d'antécédent ORL."
    assert tr.canonique(brut) == brut and tr.en_clair(brut) == brut


def test_en_clair_sans_numerotation_pour_les_verificateurs():
    md = "Le score **-2,1 ET** :\n1. phono\n2. lecture"
    assert tr.en_clair(md) == "Le score -2,1 ET :\n\n1. phono\n2. lecture"
    assert tr.en_clair(md, numeroter=False) == "Le score -2,1 ET :\n\n- phono\n- lecture"
    assert tr.contient_mise_en_forme(md) and not tr.contient_mise_en_forme("Texte.")


def test_ligne_moins_deux_et_n_est_pas_une_puce():
    """« - 2 ET » en tête de ligne est un écart-type, pas une puce : la forme
    canonique le protège, l'analyse le restitue, la version en clair l'affiche."""
    blocs = [tr.Paragraphe([tr.Segment("Écart :\n- 2 ET")])]
    md = tr.serialiser(blocs)
    assert md == "Écart :\n\\- 2 ET"
    assert tr.en_clair(md) == "Écart :\n- 2 ET"
    assert isinstance(tr.analyser(md)[0], tr.Paragraphe) and len(tr.analyser(md)) == 1


# --- Exports ---------------------------------------------------------------------

BILAN = {
    "id": 7, "statut": "valide", "type": "initial_simple", "domaine_titres": "Langage écrit",
    "epreuves": [], "cotation": None,
    "sections": [{"cle": "anamnese", "titre": "Anamnèse", "contenu":
        "Le test **Alouette** est *chuté*.\n<u>Compréhension</u> : bonne.\n\n"
        "Axes :\n- phono\n- lecture\n  suite\n\n1. un\n2. deux\n\nAutre :\n1. alpha\n2. beta"}],
}


def _runs(doc):
    return [(p.style.name, [(r.text, r.bold, r.italic, r.underline) for r in p.runs])
            for p in doc.paragraphs if p.text.strip()]


def test_export_docx_runs_et_listes():
    data = export.to_docx(BILAN)
    doc = Document(io.BytesIO(data))
    styles = _runs(doc)
    normal = next(runs for nom, runs in styles if runs and runs[0][0] == "Le test ")
    assert ("Alouette", True, None, None) in normal
    assert ("chuté", None, True, None) in normal
    assert ("Compréhension", None, None, True) in normal
    assert [nom for nom, _ in styles].count("List Bullet") == 2
    assert [nom for nom, _ in styles].count("List Number") == 4
    # L'élément « lecture » garde sa ligne de continuation (saut de ligne Word).
    lecture = next(runs for nom, runs in styles if runs and runs[0][0] == "lecture")
    assert "".join(r[0] for r in lecture) == "lecture\nsuite"
    # Deux listes numérotées = deux numérotations distinctes (chacune repart à 1).
    xml = zipfile.ZipFile(io.BytesIO(data)).read("word/document.xml").decode()
    num_ids = re.findall(r'<w:numId w:val="(\d+)"/>', xml)
    assert len(set(num_ids)) == 2 and len(num_ids) == 4


def test_export_pdf_md_txt():
    pdf = export.to_pdf(BILAN)
    assert pdf.startswith(b"%PDF")
    from pypdf import PdfReader
    texte = PdfReader(io.BytesIO(pdf)).pages[0].extract_text()
    assert "phono" in texte and "alpha" in texte and "**" not in texte
    md = export.to_markdown(BILAN)
    assert "**Alouette**" in md and "\n- phono\n- lecture\n  suite\n" in md and "\n1. alpha\n2. beta" in md
    txt = export.to_txt(BILAN)
    assert "Le test Alouette est chuté." in txt and "**" not in txt and "- phono" in txt


def test_export_ne_laisse_passer_aucune_balise_du_texte():
    """Une balise présente dans le texte (le modèle en produit parfois) reste du
    texte dans le PDF : seules NOS balises de mise en forme passent."""
    b = dict(BILAN, sections=[{"cle": "a", "titre": "A", "contenu": "<b>faux</b> & <para>x</para> **vrai**"}])
    assert export.to_pdf(b).startswith(b"%PDF")
    doc = Document(io.BytesIO(export.to_docx(b)))
    assert any("<b>faux</b>" in p.text for p in doc.paragraphs)


# --- Import Word : la mise en forme du praticien entre dans les extraits -----------

def _docx(paragraphes) -> bytes:
    d = Document()
    for p in paragraphes:
        p(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_import_docx_conserve_gras_souligne_et_listes():
    def titre(d):
        d.add_paragraph("ANAMNÈSE", style="Heading 1")

    def prose(d):
        p = d.add_paragraph()
        p.add_run("Le test ")
        p.add_run("Alouette ").bold = True
        p.add_run("est chuté chez ")
        p.add_run("DUPONT Léa").bold = True

    def souligne(d):
        p = d.add_paragraph()
        p.add_run("Compréhension").underline = True
        p.add_run(" : bonne.")

    def titre_gras(d):
        d.add_paragraph("EPREUVES").runs[0].bold = True

    def listes(d):
        d.add_paragraph("phonologie", style="List Bullet")
        d.add_paragraph("lecture", style="List Bullet")
        d.add_paragraph("un", style="List Number")

    texte = importer.extract_text(_docx([titre, prose, souligne, titre_gras, listes]), "b.docx")
    assert texte.splitlines() == [
        "ANAMNÈSE",
        "Le test **Alouette** est chuté chez **DUPONT Léa**",
        "<u>Compréhension</u> : bonne.",
        "**EPREUVES**",
        "- phonologie",
        "- lecture",
        "1. un",
    ]
    # Le titre en gras du praticien est reconnu comme en-tête de rubrique, en clair.
    sections = importer.sectionize(texte)
    assert [(c, t) for c, t, _ in sections] == [("anamnese", "ANAMNÈSE"), ("epreuves", "EPREUVES")]
    assert sections[1][2] == "- phonologie\n- lecture\n1. un"
    # L'anonymisation traverse les marqueurs.
    caviarde, n = anonymisation.caviarder(texte)
    assert "**[NOM] Léa**" in caviarde and n >= 1


# --- Chaîne IA : consigne, vérificateurs, réglage -----------------------------------

def test_consigne_de_mise_en_forme_suit_le_reglage():
    args = dict(transcription="x", sections=[], domaine_titres="", guidance="", tests_connus="")
    u = prompts.build_structure_user(**args, style_prefs={"vouvoiement": True})
    assert prompts.MISE_EN_FORME_AUTORISEE in u
    u = prompts.build_structure_user(**args, style_prefs={"mise_en_forme_ia": False})
    assert prompts.MISE_EN_FORME_INTERDITE in u and prompts.MISE_EN_FORME_AUTORISEE not in u


def test_prefixe_de_titre_souligne_ou_gras_retire():
    assert bilan.nettoyer_prefixe_titre("<u>Anamnèse</u> : né à terme.", "Anamnèse") == "né à terme."
    assert bilan.nettoyer_prefixe_titre("**Anamnèse :** né à terme.", "Anamnèse") == "né à terme."
    assert bilan.nettoyer_prefixe_titre("**Antécédents** : aucun.", "Anamnèse") == "**Antécédents** : aucun."


def test_structure_verifie_en_clair_et_respecte_le_reglage(client, monkeypatch):
    """Un score en gras reste retrouvé dans la dictée (pas de faux signalement) ;
    un numéro de liste n'est pas un chiffre à retrouver ; réglage désactivé,
    le texte du modèle est remis en clair avant d'entrer dans la rubrique."""
    reponse = (
        '{"updates":[{"section":"projet","texte":"Score **-2,1 ET** à l\'Alouette.\\n'
        '1. phonologie\\n2. lecture"}],"questions":[]}'
    )

    async def fake_chat_json(system, user, **kw):
        return reponse

    monkeypatch.setattr(llm, "chat_json", fake_chat_json)
    bid = client.post("/api/bilans", json={"domaines": ["langage_ecrit"]}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure",
                    json={"transcription": "Alouette moins deux virgule un écart-type, axes phonologie et lecture."})
    assert r.status_code == 200
    res = r.json()
    projet = next(s for s in res["bilan"]["sections"] if s["cle"] == "projet")
    assert projet["contenu"] == "Score **-2,1 ET** à l'Alouette.\n1. phonologie\n2. lecture"
    assert not [c for c in res.get("chiffres_a_verifier", []) if c["section"] == "projet"], res

    client.put("/api/config", json={"overrides": {"style": {"mise_en_forme_ia": False}}})
    bid2 = client.post("/api/bilans", json={"domaines": ["langage_ecrit"]}).json()["id"]
    res = client.post(f"/api/bilans/{bid2}/structure",
                      json={"transcription": "Alouette moins deux virgule un écart-type, axes phonologie et lecture."}).json()
    projet = next(s for s in res["bilan"]["sections"] if s["cle"] == "projet")
    assert projet["contenu"] == "Score -2,1 ET à l'Alouette.\n\n1. phonologie\n2. lecture"
