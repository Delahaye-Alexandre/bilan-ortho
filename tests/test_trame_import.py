"""Reprendre la trame d'un bilan importé (lot C du plan « mise en forme ») :
lignes structurées (titres stylés, gras), trame proposée d'après les
intitulés, découpage des extraits par titre stylé, recherche par intitulé,
route d'analyse et proposition renvoyée à l'import."""
import io

from docx import Document

from app import importer, rag, security
from tests.conftest import fake_vec
from tests.test_unites import _odt_minimal


def docx(paragraphes: list) -> bytes:
    """Un .docx à partir de (texte, style) ou (texte, style, gras)."""
    d = Document()
    for p in paragraphes:
        texte, style, gras = (p + (False,))[:3]
        par = d.add_paragraph(style=style) if style else d.add_paragraph()
        par.add_run(texte).bold = gras or None
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


PROSE = "Phrase de contenu suffisamment longue pour ne pas passer pour un intertitre."

DOCX_STYLES = docx([
    ("Compte rendu de bilan orthophonique", "Title"),
    ("Motif de la demande", "Heading 1"), (PROSE, None),
    ("Anamnèse", "Heading 1"), (PROSE, None),
    ("Contexte scolaire", "Heading 1"), (PROSE, None),
    ("Bilan analytique", "Heading 1"),
    ("Lecture", "Heading 2"), (PROSE, None),
    ("Orthographe", "Heading 2"), (PROSE, None),
    ("Remarque", None, True), (PROSE, None),   # gras : ignoré, les styles font foi
    ("Conclusion", "Heading 1"), (PROSE, None),
    ("Projet thérapeutique", "Heading 1"), (PROSE, None),
])

TEXTE_NU = """Compte rendu de bilan orthophonique
Anamnèse
Enfant né à terme, marche à 12 mois, premiers mots vers 18 mois selon la famille.
Contexte scolaire
Scolarisé en CE1, difficultés signalées par l'enseignante depuis la rentrée dernière.
Épreuves et résultats
Alouette-R : 5e percentile.
Marche à 12 mois
Conclusion
Les résultats orientent vers une évaluation complémentaire, à confirmer dans six mois.
"""


# --- lignes structurées ----------------------------------------------------------

def test_lignes_docx_portent_niveau_et_gras():
    lignes = importer.extraire_lignes(docx([
        ("Anamnèse", "Heading 1"), ("Contexte", None, True), (PROSE, None),
        ("phonologie", "List Bullet", True),
    ]), "b.docx")
    assert [(ligne.texte, ligne.niveau, ligne.gras) for ligne in lignes] == [
        ("Anamnèse", 1, False), ("**Contexte**", None, True), (PROSE, None, False),
        ("- **phonologie**", None, False),   # un élément de liste n'est pas un intertitre
    ]
    # Le texte brut reste celui d'avant (extraits, anonymisation, tests existants).
    assert importer.extract_text(docx([("Anamnèse", "Heading 1"), (PROSE, None)]), "b.docx") \
        == f"Anamnèse\n{PROSE}"
    assert importer._niveau_style("Titre 2") == 2 and importer._niveau_style("Title") == 0
    assert importer._niveau_style("Sous-titre") is None and importer._niveau_style("Normal") is None


def test_lignes_odt_niveau_de_plan():
    data = _odt_minimal(
        '<text:h text:outline-level="1">Anamnèse</text:h>'
        f"<text:p>{PROSE}</text:p>"
        '<text:h text:outline-level="2">Contexte familial</text:h>'
        f"<text:p>{PROSE}</text:p>"
        "<text:h>Projet thérapeutique</text:h>"
    )
    lignes = importer.extraire_lignes(data, "b.odt")
    assert [(ligne.texte, ligne.niveau) for ligne in lignes] == [
        ("Anamnèse", 1), (PROSE, None), ("Contexte familial", 2), (PROSE, None),
        ("Projet thérapeutique", 1),
    ]


def test_cle_de_rubrique():
    assert importer.cle_de_rubrique("Contexte scolaire") == "contexte_scolaire"
    assert importer.cle_de_rubrique("Épreuves & résultats — 2024 :") == "epreuves_resultats_2024"
    assert importer.cle_de_rubrique("   ") == "rubrique"
    assert len(importer.cle_de_rubrique("mot " * 30)) <= 40
    assert importer._titre_propre("1. **Anamnèse** :") == "Anamnèse"
    assert importer._titre_propre("II - Épreuves   et   résultats") == "Épreuves et résultats"


# --- trame proposée ---------------------------------------------------------------

def test_trame_proposee_depuis_les_styles_word():
    p = importer.proposer_trame(DOCX_STYLES, "bilan.docx")
    assert p["detection"] == "styles"
    assert p["sections"] == [
        {"cle": "motif_de_la_demande", "titre": "Motif de la demande"},
        {"cle": "anamnese", "titre": "Anamnèse"},
        {"cle": "contexte_scolaire", "titre": "Contexte scolaire"},
        {"cle": "epreuves", "titre": "Bilan analytique"},      # mot-clé : clé du tronc commun
        {"cle": "diagnostic", "titre": "Conclusion"},
        {"cle": "projet", "titre": "Projet thérapeutique"},
    ]
    # Ni le titre du document, ni les sous-titres (niveau 2), ni le gras.
    titres = [s["titre"] for s in p["sections"]]
    assert not {"Compte rendu de bilan orthophonique", "Lecture", "Orthographe", "Remarque"} & set(titres)


def test_trame_proposee_depuis_le_gras_word():
    p = importer.proposer_trame(docx([
        ("1. Anamnèse :", None, True), (PROSE, None),
        ("Contexte familial", None, True), (PROSE, None),
        ("Le bilan a été passé en deux séances d'une heure, à une semaine d'intervalle.", None, True),
        (PROSE, None),
        ("Épreuves", None, True), ("- phonologie", None), ("- lecture", None),
    ]), "bilan.docx")
    assert p["detection"] == "gras"
    assert p["sections"] == [
        {"cle": "anamnese", "titre": "Anamnèse"},
        {"cle": "contexte_familial", "titre": "Contexte familial"},
        {"cle": "epreuves", "titre": "Épreuves"},
    ]


def test_trame_proposee_depuis_un_texte_nu():
    p = importer.proposer_trame(TEXTE_NU.encode(), "bilan.txt")
    assert p["detection"] == "lignes"
    assert [s["cle"] for s in p["sections"]] == ["anamnese", "contexte_scolaire", "epreuves", "diagnostic"]
    # « Marche à 12 mois » (chiffre) et le titre du document ne sont pas des rubriques.
    assert "Marche à 12 mois" not in [s["titre"] for s in p["sections"]]


def test_trame_proposee_rien_de_net():
    assert importer.proposer_trame(b"Un texte sans aucun titre reconnaissable, tout en prose.", "n.txt") is None
    assert importer.proposer_trame(b"Anamnese\nSeule rubrique.", "n.txt") is None
    # Une liste de trente lignes courtes n'est pas une trame.
    bruit = "\n".join(f"Titre {chr(65 + i)}bis\n{PROSE}" for i in range(30))
    assert importer.proposer_trame(bruit.encode(), "n.txt") is None


def test_trame_proposee_cles_uniques():
    p = importer.proposer_trame(docx([
        ("Observations", "Heading 1"), (PROSE, None),
        ("Épreuves", "Heading 1"), (PROSE, None),
        ("Résultats", "Heading 1"), (PROSE, None),       # même mot-clé : clé déduite du titre
        ("Observations", "Heading 1"), (PROSE, None),    # doublon : ignoré
    ]), "b.docx")
    assert [s["cle"] for s in p["sections"]] == ["observations", "epreuves", "resultats"]


# --- découpage des extraits et recherche par intitulé -------------------------------

def test_decoupage_par_titres_styles_garde_la_cle_courante():
    chunks = importer.decouper(DOCX_STYLES, "bilan.docx")
    assert [(c, t) for c, t, _ in chunks] == [
        ("global", "Extrait"),                    # titre du document, avant tout en-tête
        ("global", "Motif de la demande"),        # avant le premier mot-clé : écarté à l'import
        ("anamnese", "Anamnèse"),
        ("anamnese", "Contexte scolaire"),        # titre sans mot-clé : clé héritée
        ("epreuves", "Lecture"), ("epreuves", "Orthographe"),
        ("diagnostic", "Conclusion"), ("projet", "Projet thérapeutique"),
    ]
    # Le gras n'est pas un en-tête à l'import (extraits moins morcelés) :
    # « Remarque » reste dans l'extrait « Orthographe ».
    assert "**Remarque**" in chunks[5][2]
    assert all(contenu.startswith(PROSE) for _, _, contenu in chunks[2:])
    # Un texte nu se découpe comme avant (mots-clés seulement).
    assert [c for c, _, _ in importer.decouper(TEXTE_NU.encode(), "b.txt")] == [
        "global", "anamnese", "epreuves", "diagnostic",
    ]


def test_retrieve_retrouve_un_extrait_par_son_intitule(con):
    rag.add_reference(con, None, "import", "langage_oral", "anamnese", "Contexte scolaire",
                      "Scolarisé en CE1.", fake_vec("scolaire"))
    rag.add_reference(con, None, "import", "langage_oral", "anamnese", "Anamnèse",
                      "Né à terme.", fake_vec("naissance"))
    par_intitule = rag.retrieve(con, fake_vec("scolaire"), section_cle="contexte_scolaire")
    assert [r["titre"] for r in par_intitule] == ["Contexte scolaire"]
    par_cle = rag.retrieve(con, fake_vec("scolaire"), section_cle="anamnese")
    assert {r["titre"] for r in par_cle} == {"Contexte scolaire", "Anamnèse"}
    assert rag.retrieve(con, fake_vec("scolaire"), section_cle="projet") == []


# --- API ---------------------------------------------------------------------------

def test_api_analyse_trame(client):
    r = client.post("/api/config/trame/analyse",
                    files={"fichier": ("bilan-DUPONT.docx", DOCX_STYLES, "application/octet-stream")})
    assert r.status_code == 200
    assert r.json()["detection"] == "styles"
    assert [s["cle"] for s in r.json()["sections"]][:3] == ["motif_de_la_demande", "anamnese", "contexte_scolaire"]
    # Rien d'enregistré : la trame effective reste la trame réglementaire.
    assert "trame" not in client.get("/api/config/overrides").json()
    r = client.post("/api/config/trame/analyse",
                    files={"fichier": ("prose.txt", b"De la prose, sans aucun titre.", "text/plain")})
    assert r.status_code == 422 and "Moins de deux rubriques" in r.json()["detail"]
    assert client.post("/api/config/trame/analyse",
                       files={"fichier": ("vide.txt", b"", "text/plain")}).status_code == 400
    with security.transaction() as con:
        details = [r[0] for r in con.execute(
            "SELECT details FROM audit_log WHERE action='config_trame'").fetchall()]
    assert any(d.startswith("analyse · fichier .docx") for d in details)
    assert not any("DUPONT" in d for d in details)   # le nom du fichier n'est pas journalisé
    # La proposition s'enregistre telle quelle par la route de la trame.
    assert client.put("/api/config/trame", json={"sections": r.json() and
                      client.post("/api/config/trame/analyse",
                                  files={"fichier": ("b.docx", DOCX_STYLES, "")}).json()["sections"]}
                      ).status_code == 200


def test_api_import_reference_propose_la_trame(client, mock_embed):
    r = client.post("/api/references",
                    files={"file": ("bilan.docx", DOCX_STYLES, "application/octet-stream")},
                    data={"domaine": "langage_oral"})
    assert r.status_code == 200
    p = r.json()["trame_proposee"]
    assert p["detection"] == "styles" and len(p["sections"]) == 6
    assert r.json()["n"] == 6   # les extraits, blocs avant le premier mot-clé écartés
    r = client.post("/api/references",
                    files={"file": ("prose.txt", b"De la prose, sans aucun titre.", "text/plain")},
                    data={"domaine": ""})
    assert r.status_code == 200 and r.json()["trame_proposee"] is None
