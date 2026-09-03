"""Mise en page des exports (lot B du plan « mise en forme ») : réglages
`mise_en_page` appliqués au Word et au PDF, logo, numérotation, numéros de
page, bilan d'exemple, routes logo et aperçu."""
import base64
import io
import re
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image
from pypdf import PdfReader

from app import config, export

BILAN = {
    "id": 1, "type": "initial_simple", "statut": "valide", "domaine_titres": "Voix",
    "date_bilan": "2026-07-25", "created_at": "2026-07-25 10:00:00",
    "sections": [
        {"titre": "Anamnèse", "contenu": "Né à terme. **Audition** contrôlée."},
        {"titre": "Projet thérapeutique", "contenu": "- une\n- deux"},
    ],
    "epreuves": [{"test_nom": "EXALANG", "resultats": [
        {"sous_epreuve": "Dictée", "score_brut": "14/30", "etalonnage_type": "ecart_type",
         "etalonnage_valeur": "-2,0", "drapeau_seuil": "severe"},
    ]}],
}


def cfg_avec(**mep) -> dict:
    return config._deep_merge(config.DEFAULTS, {"mise_en_page": mep})


def png(largeur: int, hauteur: int, mode: str = "RGBA") -> bytes:
    buf = io.BytesIO()
    Image.new(mode, (largeur, hauteur), (30, 60, 120, 255) if mode == "RGBA" else 128).save(buf, "PNG")
    return buf.getvalue()


def jpeg(largeur: int, hauteur: int) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (largeur, hauteur), (200, 30, 30)).save(buf, "JPEG")
    return buf.getvalue()


def docx_xml(data: bytes, membre: str = "word/document.xml") -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(membre).decode() if membre in z.namelist() else ""


# --- réglages effectifs -------------------------------------------------------

def test_mise_en_page_effective_complete_les_defauts():
    assert export.mise_en_page(None) == config.DEFAULTS["mise_en_page"]
    mp = export.mise_en_page({"mise_en_page": {"police": "Georgia"}})
    assert mp["police"] == "Georgia" and mp["marges_mm"] == 20
    # Une config sans la section (ancienne surcharge, tests) : les défauts.
    assert export.mise_en_page({"praticien": {}})["taille_corps"] == 11


# --- Word ---------------------------------------------------------------------

def test_docx_police_taille_interligne_marges_titres_et_pied_de_page():
    data = export.to_docx(BILAN, cfg_avec(
        police="Georgia", taille_corps=12, interligne=1.5, marges_mm=25,
        couleur_titres="#1f3a5f", numeros_de_page=True,
    ))
    doc = Document(io.BytesIO(data))
    normal = doc.styles["Normal"]
    assert normal.font.name == "Georgia" and normal.font.size == Pt(12)
    assert normal.paragraph_format.line_spacing == 1.5
    rfonts = normal.element.rPr.rFonts
    assert rfonts.get(qn("w:eastAsia")) == "Georgia" and rfonts.get(qn("w:cs")) == "Georgia"
    for nom, delta in (("Heading 1", 6), ("Heading 2", 2)):
        st = doc.styles[nom]
        assert st.font.name == "Georgia" and st.font.size == Pt(12 + delta)
        assert str(st.font.color.rgb) == "1F3A5F"
        # Les renvois au thème (Calibri Light, bleu accent) sont retirés :
        # sinon Word les préfère au nom et à la couleur posés.
        assert st.element.rPr.rFonts.get(qn("w:asciiTheme")) is None
        assert st.element.rPr.color.get(qn("w:themeColor")) is None
    s = doc.sections[0]
    # Word range les marges en twips : à l'arrondi près.
    assert [round(m.mm) for m in (s.left_margin, s.right_margin, s.top_margin, s.bottom_margin)] == [25] * 4
    pied = docx_xml(data, "word/footer1.xml")
    assert 'w:instr="PAGE"' in pied and 'w:instr="NUMPAGES"' in pied and "Page " in pied
    # Titres du document : Heading 1 puis Heading 2 (relus par importer.py).
    xml = docx_xml(data)
    assert 'w:val="Heading1"' in xml and 'w:val="Heading2"' in xml


def test_docx_sans_numeros_de_page_ni_logo():
    data = export.to_docx(BILAN, cfg_avec(numeros_de_page=False))
    assert "NUMPAGES" not in docx_xml(data, "word/footer1.xml")
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert not [n for n in z.namelist() if n.startswith("word/media/")]


# --- numérotation des rubriques ----------------------------------------------

def test_rubriques_numerotees_dans_les_quatre_formats():
    cfg = cfg_avec(rubriques_numerotees=True)
    md = export.to_markdown(BILAN, cfg)
    assert "## 1. Anamnèse" in md and "## 2. Projet thérapeutique" in md
    assert "## 3. Résultats des épreuves" in md  # numérotés à la suite, pas seulement les rubriques
    txt = export.to_txt(BILAN, cfg)
    assert "1. ANAMNÈSE" in txt and "3. RÉSULTATS DES ÉPREUVES" in txt
    assert "1. Anamnèse" in docx_xml(export.to_docx(BILAN, cfg))
    assert export.to_pdf(BILAN, cfg).startswith(b"%PDF")
    # Par défaut : rien de numéroté.
    assert "## 1." not in export.to_markdown(BILAN, config.DEFAULTS)


# --- logo ---------------------------------------------------------------------

def test_preparer_logo_verifie_reduit_et_reencode():
    logo = export.preparer_logo(png(800, 1000))
    assert logo["type"] == "image/png" and (logo["largeur"], logo["hauteur"]) == (320, 400)
    assert Image.open(io.BytesIO(base64.b64decode(logo["donnees"]))).mode == "RGBA"
    petit = export.preparer_logo(png(120, 60, mode="L"))
    assert (petit["largeur"], petit["hauteur"]) == (120, 60)  # pas agrandi
    assert export.preparer_logo(jpeg(300, 900))["type"] == "image/jpeg"
    with pytest.raises(ValueError, match="illisible"):
        export.preparer_logo(b"pas une image du tout")
    gif = io.BytesIO()
    Image.new("P", (10, 10)).save(gif, "GIF")
    with pytest.raises(ValueError, match="GIF"):
        export.preparer_logo(gif.getvalue())


def test_logo_dans_le_word_et_le_pdf_pas_dans_le_markdown():
    cfg = cfg_avec(logo=export.preparer_logo(png(400, 200)), logo_position="droite",
                   logo_hauteur_mm=25)
    data = export.to_docx(BILAN, cfg)
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert [n for n in z.namelist() if n.startswith("word/media/")]
    doc = Document(io.BytesIO(data))
    avec_image = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
    assert len(avec_image) == 1 and avec_image[0].alignment == 2  # RIGHT
    pdf = export.to_pdf(BILAN, cfg)
    assert pdf.startswith(b"%PDF") and b"/Image" in pdf
    md = export.to_markdown(BILAN, cfg)
    assert "logo" not in md.lower() and "Compte-rendu" in md


def test_logo_illisible_n_empeche_pas_l_export():
    cfg = cfg_avec(logo={"type": "image/png",
                         "donnees": base64.b64encode(b"pas une image").decode()})
    assert export.to_docx(BILAN, cfg)
    assert export.to_pdf(BILAN, cfg).startswith(b"%PDF")
    cfg = cfg_avec(logo={"type": "image/png", "donnees": "%%pas du base64%%"})
    assert export.to_pdf(BILAN, cfg).startswith(b"%PDF")


# --- PDF ----------------------------------------------------------------------

def test_pdf_numeros_de_page_sur_plusieurs_pages():
    long = {**BILAN, "sections": [
        {"titre": f"Rubrique {i}", "contenu": "Phrase de remplissage. " * 60} for i in range(12)
    ]}
    def pieds(pdf: bytes) -> list[list[str]]:
        # Les flux de page sont compressés : relire le texte avec pypdf.
        return [re.findall(r"Page \d+ / \d+", p.extract_text())
                for p in PdfReader(io.BytesIO(pdf)).pages]

    numerotes = pieds(export.to_pdf(long, cfg_avec(numeros_de_page=True)))
    n = len(numerotes)
    assert n >= 2 and numerotes == [[f"Page {i} / {n}"] for i in range(1, n + 1)]
    assert all(p == [] for p in pieds(export.to_pdf(long, cfg_avec(numeros_de_page=False))))
    # Brouillon : le filigrane et le pied de page cohabitent.
    brouillon = PdfReader(io.BytesIO(export.to_pdf({**long, "statut": "brouillon"}, cfg_avec())))
    assert "BROUILLON" in brouillon.pages[1].extract_text()
    assert "Page 2 /" in brouillon.pages[1].extract_text()


def test_pdf_marges_larges_et_tableau_qui_tient():
    """Des largeurs de colonnes fixes débordaient de la page dès que les marges
    s'élargissaient : elles suivent maintenant la largeur utile."""
    pdf = export.to_pdf(BILAN, cfg_avec(marges_mm=40, taille_corps=16))
    assert pdf.startswith(b"%PDF")


def test_pdf_polices_de_repli_et_truetype(monkeypatch):
    monkeypatch.setattr(export, "_POLICES_PDF", {})
    monkeypatch.setattr(export, "_CHEMINS_POLICES", {})
    monkeypatch.setattr(export, "_trouver_police", lambda nom: None)
    assert export._polices_pdf("Times New Roman")["normal"] == "Times-Roman"
    assert export._polices_pdf("Georgia")["gras"] == "Times-Bold"
    assert export._polices_pdf("Arial")["italique"] == "Helvetica-Oblique"
    assert export._polices_pdf("Police inconnue")["normal"] == "Helvetica"
    # Fichier TrueType trouvé (ici la Vera livrée avec reportlab, en guise
    # d'Arial) : la famille est enregistrée, incorporée au PDF, et une face
    # manquante retombe sur la régulière.
    from pathlib import Path

    import reportlab

    vera = Path(reportlab.__file__).parent / "fonts" / "Vera.ttf"
    monkeypatch.setattr(export, "_POLICES_PDF", {})
    monkeypatch.setattr(export, "_CHEMINS_POLICES", {})
    monkeypatch.setattr(export, "_trouver_police", lambda nom: vera if nom == "arial.ttf" else None)
    polices = export._polices_pdf("Arial")
    assert polices["normal"] == "MEP-arial" and polices["gras"] == "MEP-arial-Bold"
    pdf = export.to_pdf(BILAN, cfg_avec(police="Arial", numeros_de_page=False))
    assert b"BitstreamVeraSans" in pdf


# --- bilan d'exemple ----------------------------------------------------------

def test_bilan_exemple_passe_par_tous_les_blocs():
    cfg = config._deep_merge(config.DEFAULTS, {
        "praticien": {"nom": "Martin", "prenom": "Claire"},
        "cotation": {"valeur_amo": 3.0, "bilan_simple_coeff": 20},
    })
    b = export.bilan_exemple(cfg)
    kinds = [k for k, _ in export._content(b, cfg)]
    for attendu in ("entete", "dest", "h1", "riche", "table", "sign", "i"):
        assert attendu in kinds
    assert "brouillon" not in kinds
    assert b["cotation"] == {"coefficient": 20.0, "montant": 60.0, "valeur_lettre_cle": 3.0}
    md = export.to_markdown(b, cfg)
    assert "EXEMPLE Camille" in md and "Claire Martin" in md and "60,00" in md
    assert export.to_pdf(b, cfg).startswith(b"%PDF") and export.to_docx(b, cfg)


# --- API ----------------------------------------------------------------------

def test_api_reglages_de_mise_en_page(client):
    eff = client.put("/api/config", json={"overrides": {"mise_en_page": {
        "police": "Georgia", "marges_mm": 25, "rubriques_numerotees": True,
    }}}).json()
    assert eff["mise_en_page"]["police"] == "Georgia" and eff["mise_en_page"]["marges_mm"] == 25
    assert eff["mise_en_page"]["taille_corps"] == 11  # fusion : le reste garde ses défauts
    for mauvais in ({"couleur_titres": "rouge"}, {"taille_corps": 40}, {"logo_position": "haut"},
                    {"marges_mm": 90}, {"logo": {"donnees": "x"}}):
        r = client.put("/api/config", json={"overrides": {"mise_en_page": mauvais}})
        assert r.status_code == 422, mauvais
    # Retour aux valeurs recommandées, clé par clé ou en bloc.
    eff = client.delete("/api/config/mise_en_page", params={"cles": "police"}).json()
    assert eff["mise_en_page"]["police"] == "Calibri" and eff["mise_en_page"]["marges_mm"] == 25
    eff = client.delete("/api/config/mise_en_page").json()
    assert eff["mise_en_page"] == config.DEFAULTS["mise_en_page"]


def test_api_logo_depot_retrait_et_export(client):
    r = client.put("/api/config/logo", files={"fichier": ("logo.png", png(600, 900), "image/png")})
    assert r.status_code == 200
    logo = r.json()["mise_en_page"]["logo"]
    assert logo["type"] == "image/png" and logo["hauteur"] == 400 and logo["largeur"] == 267
    assert client.get("/api/config/overrides").json()["mise_en_page"]["logo"]["hauteur"] == 400
    # L'export d'un bilan porte l'image (Word) et le PDF se construit.
    bid = client.post("/api/bilans", json={"domaines": ["voix"]}).json()["id"]
    docx = client.get(f"/api/bilans/{bid}/export", params={"format": "docx"}).content
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        assert [n for n in z.namelist() if n.startswith("word/media/")]
    assert client.get(f"/api/bilans/{bid}/export", params={"format": "pdf"}).content.startswith(b"%PDF")
    # Fichiers refusés : texte déguisé, format non pris en charge, trop gros.
    r = client.put("/api/config/logo", files={"fichier": ("logo.png", b"<svg/>", "image/png")})
    assert r.status_code == 422 and "PNG ou JPEG" in r.json()["detail"]
    gif = io.BytesIO()
    Image.new("P", (10, 10)).save(gif, "GIF")
    assert client.put("/api/config/logo", files={"fichier": ("l.gif", gif.getvalue(), "image/gif")}).status_code == 422
    trop = b"\x89PNG" + b"0" * export.TAILLE_MAX_LOGO
    assert client.put("/api/config/logo", files={"fichier": ("l.png", trop, "image/png")}).status_code == 413
    # Un refus ne touche pas le logo en place ; le retrait garde les autres réglages.
    client.put("/api/config", json={"overrides": {"mise_en_page": {"police": "Arial"}}})
    assert client.get("/api/config").json()["mise_en_page"]["logo"]["hauteur"] == 400
    eff = client.delete("/api/config/logo").json()
    assert eff["mise_en_page"]["logo"] is None and eff["mise_en_page"]["police"] == "Arial"
    assert "logo" not in client.get("/api/config/overrides").json()["mise_en_page"]


def test_api_apercu_de_mise_en_page(client):
    r = client.post("/api/config/mise_en_page/apercu", json={"police": "Georgia", "marges_mm": 15})
    assert r.status_code == 200 and r.headers["content-type"] == "application/pdf"
    assert r.content.startswith(b"%PDF") and "inline" in r.headers["content-disposition"]
    assert client.post("/api/config/mise_en_page/apercu", json={"taille_corps": 2}).status_code == 422
    # Le logo enregistré entre dans l'aperçu même si l'écran ne l'envoie pas.
    client.put("/api/config/logo", files={"fichier": ("logo.png", png(200, 100), "image/png")})
    assert b"/Image" in client.post("/api/config/mise_en_page/apercu", json={}).content
    # Rien n'a été enregistré par les aperçus.
    assert client.get("/api/config").json()["mise_en_page"]["police"] == "Calibri"
