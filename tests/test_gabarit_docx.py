"""Gabarit Word du praticien (lot D du plan « mise en forme ») : dépôt vérifié
et décrit, export Word qui part du gabarit (en-tête, pied de page, sections,
styles gardés, corps vidé), replis pour les styles absents, routes."""
import io
import zipfile
from datetime import date

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from app import config, export, security

BILAN = {
    "id": 1, "type": "initial_simple", "statut": "valide", "domaine_titres": "Voix",
    "date_bilan": "2026-07-25", "created_at": "2026-07-25 10:00:00",
    "sections": [
        {"titre": "Anamnèse", "contenu": "Né à terme. **Audition** contrôlée."},
        {"titre": "Projet thérapeutique", "contenu": "- une\n- deux\n\n1. alpha\n2. beta"},
    ],
    "epreuves": [{"test_nom": "EXALANG", "resultats": [
        {"sous_epreuve": "Dictée", "score_brut": "14/30", "etalonnage_type": "ecart_type",
         "etalonnage_valeur": "-2,0", "drapeau_seuil": "severe"},
    ]}],
}
PRATICIEN = {"nom": "Martin", "prenom": "Camille", "titre": "Orthophoniste",
             "adresse": "12 rue des Lilas", "code_postal": "59000", "ville": "Lille"}
EN_TETE = "Cabinet d'orthophonie Exemple — 12 rue des Lilas"


def cfg_avec(**mep) -> dict:
    return config._deep_merge(config.DEFAULTS, {"praticien": PRATICIEN, "mise_en_page": mep})


def gabarit(en_tete: str | None = EN_TETE, pied: str | None = None, sans: tuple = (),
            corps: tuple = ("Votre texte ici", "Deuxième paragraphe de remplissage")) -> bytes:
    """Gabarit de test : police et taille propres, Titre 1 rouge sombre, marges
    de 30 mm, en-tête (et pied) garnis ; `sans` retire des styles, comme un
    papier à en-tête qui n'a jamais servi à écrire un titre ou une liste."""
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Mm(30)
    if en_tete:
        s.header.paragraphs[0].text = en_tete
    if pied:
        s.footer.paragraphs[0].text = pied
    doc.styles["Normal"].font.name = "Garamond"
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(0x88, 0x00, 0x00)
    for texte in corps:
        doc.add_paragraph(texte)
    for nom in sans:
        doc.styles[nom].delete()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def zip_remplace(data: bytes, nom_membre: str, ancien: str, nouveau: str) -> bytes:
    """Même .docx, un membre réécrit (type de contenu, XML cassé…)."""
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as z, zipfile.ZipFile(out, "w") as w:
        for info in z.infolist():
            contenu = z.read(info)
            if info.filename == nom_membre:
                assert ancien.encode() in contenu, nom_membre
                contenu = contenu.replace(ancien.encode(), nouveau.encode())
            w.writestr(info, contenu)
    return out.getvalue()


def zip_avec_types(data: bytes, ancien: str, nouveau: str) -> bytes:
    """Même .docx, type de contenu principal réécrit (modèle, macros…)."""
    return zip_remplace(data, "[Content_Types].xml", ancien, nouveau)


def en_dotx(data: bytes) -> bytes:
    return zip_avec_types(data, export._CT_DOCUMENT, export._CT_MODELE)


def membre(data: bytes, nom: str) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(nom).decode() if nom in z.namelist() else ""


def textes(doc) -> list[str]:
    return [p.text for p in doc.paragraphs]


# --- préparation et description -----------------------------------------------

def test_preparer_gabarit_decrit_le_document():
    entree = gabarit()
    data, desc = export.preparer_gabarit(entree, "Papier à en-tête (cabinet).docx")
    # Un .docx est gardé tel quel (comparé à l'entrée elle-même : deux .docx
    # générés à une seconde d'écart diffèrent par l'horodatage du zip).
    assert data == entree
    assert desc["nom"] == "Papier a en-tete cabinet.docx"
    assert desc["taille"] == len(data) and desc["depose_le"] == date.today().isoformat()
    assert desc["styles"] == list(export.STYLES_GABARIT)
    assert desc["en_tete"] is True and desc["pied_de_page"] is False
    # Pied garni, titres et listes absents : la description le dit.
    _, desc = export.preparer_gabarit(
        gabarit(pied="SIRET 000", sans=("Heading 1", "Heading 2", "List Number")), "x.docx"
    )
    assert desc["pied_de_page"] is True
    assert desc["styles"] == ["Normal", "List Bullet", "Table Grid"]
    # Nom absent ou vide : un nom par défaut, jamais un chemin.
    assert export.preparer_gabarit(gabarit(), "")[1]["nom"] == "gabarit.docx"
    assert export.preparer_gabarit(gabarit(), "C:\\Users\\moi\\..\\modèle.dotx")[1]["nom"] == "modele.docx"


def test_preparer_gabarit_convertit_un_modele_dotx():
    dotx = en_dotx(gabarit())
    with pytest.raises(ValueError):
        Document(io.BytesIO(dotx))  # python-docx n'ouvre pas un modèle
    data, desc = export.preparer_gabarit(dotx, "modele.dotx")
    assert export._CT_MODELE not in membre(data, "[Content_Types].xml")
    assert EN_TETE in membre(data, "word/header1.xml")
    assert desc["nom"] == "modele.docx" and desc["en_tete"] is True
    # Et le gabarit converti sert à l'export.
    assert EN_TETE in membre(export.to_docx(BILAN, None, data), "word/header1.xml")


def test_preparer_gabarit_refuse_ce_qui_n_est_pas_un_document_word():
    with pytest.raises(ValueError, match="document Word"):
        export.preparer_gabarit(b"Bonjour, ceci est du texte.", "bilan.txt")
    # Un zip qui n'est pas un document Word (tableur, archive quelconque).
    autre = io.BytesIO()
    with zipfile.ZipFile(autre, "w") as w:
        w.writestr("[Content_Types].xml", "<Types><Override PartName='/xl/workbook.xml' "
                   "ContentType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml'/></Types>")
    with pytest.raises(ValueError, match="document Word"):
        export.preparer_gabarit(autre.getvalue(), "classeur.xlsx")
    # Macros : refusées, que le type le dise ou qu'un projet VBA soit embarqué.
    docm = zip_avec_types(gabarit(), export._CT_DOCUMENT,
                          export._CT_DOCUMENT.replace("document.main", "document.macroEnabled.main"))
    with pytest.raises(ValueError, match="macros"):
        export.preparer_gabarit(docm, "m.docm")
    with_vba = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(gabarit())) as z, zipfile.ZipFile(with_vba, "w") as w:
        for info in z.infolist():
            w.writestr(info, z.read(info))
        w.writestr("word/vbaProject.bin", b"\x00")
    with pytest.raises(ValueError, match="macros"):
        export.preparer_gabarit(with_vba.getvalue(), "m.docx")
    # Contenu décompressé démesuré : refusé avant tout parsing.
    bombe = io.BytesIO()
    with zipfile.ZipFile(bombe, "w", zipfile.ZIP_DEFLATED) as w:
        w.writestr("[Content_Types].xml", "<Types/>")
        w.writestr("word/document.xml", b"\0" * (export.TAILLE_MAX_GABARIT_DECOMPRESSE + 1))
    assert len(bombe.getvalue()) < export.TAILLE_MAX_GABARIT
    with pytest.raises(ValueError, match="volumineux"):
        export.preparer_gabarit(bombe.getvalue(), "b.docx")
    # Document Word au zip valide mais au contenu cassé : message pour l'écran.
    casse = zip_remplace(gabarit(), "word/document.xml", "<w:body>", "<w:body")
    with pytest.raises(ValueError, match="illisible"):
        export.preparer_gabarit(casse, "c.docx")


# --- export Word sur le gabarit -------------------------------------------------

def test_docx_sur_gabarit_garde_en_tete_sections_styles_et_vide_le_corps():
    cfg = cfg_avec(police="Georgia", taille_corps=9, marges_mm=15, couleur_titres="#0000ff",
                   numeros_de_page=True)
    cfg["mise_en_page"]["logo"] = export.preparer_logo(_png())
    data = export.to_docx(BILAN, cfg, gabarit(pied="SIRET 000 — RPPS 111"))
    doc = Document(io.BytesIO(data))
    # Le corps du gabarit a disparu, le compte-rendu est là.
    assert "Votre texte ici" not in textes(doc) and "Deuxième paragraphe de remplissage" not in textes(doc)
    assert "Compte-rendu de bilan orthophonique" in textes(doc)
    assert "Anamnèse" in textes(doc)
    # En-tête et pied de page du gabarit, tels quels : pas de numéros ajoutés
    # à un pied garni.
    assert EN_TETE in membre(data, "word/header1.xml")
    pied = membre(data, "word/footer1.xml")
    assert "SIRET 000" in pied and 'w:instr="PAGE"' not in pied
    # Styles et sections du gabarit, pas les réglages de la configuration.
    normal = doc.styles["Normal"]
    assert normal.font.name == "Garamond" and normal.font.size == Pt(12)
    assert str(doc.styles["Heading 1"].font.color.rgb) == "880000"
    assert round(doc.sections[0].left_margin.mm) == 30
    # Pas de logo : le papier à en-tête est celui du gabarit.
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert not [n for n in z.namelist() if n.startswith("word/media/")]
    # L'identité du cabinet est écrite (le gabarit ne la porte pas, par défaut).
    assert "Camille Martin" in textes(doc)
    # Sans gabarit, rien ne change : les réglages s'appliquent, le logo aussi.
    sans = export.to_docx(BILAN, cfg)
    assert Document(io.BytesIO(sans)).styles["Normal"].font.name == "Georgia"


def test_docx_sur_gabarit_pied_vide_recoit_les_numeros_de_page():
    data = export.to_docx(BILAN, cfg_avec(numeros_de_page=True), gabarit())
    pied = membre(data, "word/footer1.xml")
    assert 'w:instr="PAGE"' in pied and 'w:instr="NUMPAGES"' in pied
    data = export.to_docx(BILAN, cfg_avec(numeros_de_page=False), gabarit())
    assert 'w:instr="PAGE"' not in membre(data, "word/footer1.xml")


def test_docx_sur_gabarit_identite_portee_par_le_gabarit():
    cfg = cfg_avec(gabarit_porte_identite=True)
    doc = Document(io.BytesIO(export.to_docx(BILAN, cfg, gabarit())))
    assert "Camille Martin" not in textes(doc)
    assert "12 rue des Lilas, 59000 Lille" not in textes(doc)
    # La signature reste : c'est celle du document, pas du papier.
    assert any("Martin" in t for t in textes(doc))
    # Sans gabarit, le réglage est sans effet : l'identité est écrite (Word
    # neuf comme PDF).
    assert "Camille Martin" in textes(Document(io.BytesIO(export.to_docx(BILAN, cfg))))
    assert ("entete", ["Camille Martin", "Orthophoniste", "12 rue des Lilas, 59000 Lille"]) in export._content(BILAN, cfg)


def test_docx_sur_gabarit_sans_styles_de_titre_ni_de_liste_ni_de_tableau():
    sans = ("Heading 1", "Heading 2", "List Bullet", "List Number", "Table Grid")
    cfg = cfg_avec(couleur_titres="#1f3a5f", taille_corps=8)
    data = export.to_docx(BILAN, cfg, gabarit(sans=sans))
    doc = Document(io.BytesIO(data))
    # Titres créés d'après la taille du gabarit (12) et la couleur réglée.
    for nom, delta in (("Heading 1", 6), ("Heading 2", 2)):
        st = doc.styles[nom]
        assert st.font.bold and st.font.size == Pt(12 + delta)
        assert str(st.font.color.rgb) == "1F3A5F" and st.base_style.name == "Normal"
        assert st.element.get(qn("w:customStyle")) is None  # style intégré, relu comme un titre
    titres = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert [p.text for p in titres][:2] == ["Anamnèse", "Projet thérapeutique"]
    # Listes en texte, sans paragraphe vide orphelin devant chaque élément.
    t = textes(doc)
    i = t.index("• une")
    assert t[i:i + 4] == ["• une", "• deux", "1. alpha", "2. beta"]
    # Tableau quadrillé à la main.
    assert "<w:tblBorders>" in membre(data, "word/document.xml")
    assert "<w:insideH" in membre(data, "word/document.xml")


def test_docx_sur_gabarit_taille_des_defauts_du_document():
    """Normal sans taille explicite : les titres créés partent des défauts du
    document (demi-points), puis de la configuration."""
    doc = Document()
    doc.styles["Normal"].font.size = None
    doc.styles["Heading 1"].delete()
    doc.styles.element.xpath("./w:docDefaults/w:rPrDefault/w:rPr/w:sz")[0].set(qn("w:val"), "26")
    buf = io.BytesIO()
    doc.save(buf)
    out = Document(io.BytesIO(export.to_docx(BILAN, cfg_avec(taille_corps=9), buf.getvalue())))
    assert out.styles["Heading 1"].font.size == Pt(13 + 6)


def test_docx_sur_gabarit_a_plusieurs_sections_garde_la_derniere():
    from docx.enum.section import WD_SECTION

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Première section"
    doc.add_paragraph("page de garde")
    s2 = doc.add_section(WD_SECTION.NEW_PAGE)
    s2.header.is_linked_to_previous = False
    s2.header.paragraphs[0].text = "Papier courant"
    doc.add_paragraph("corps")
    buf = io.BytesIO()
    doc.save(buf)
    data = export.to_docx(BILAN, cfg_avec(), buf.getvalue())
    out = Document(io.BytesIO(data))
    assert len(out.sections) == 1 and out.sections[0].header.paragraphs[0].text == "Papier courant"
    assert "page de garde" not in textes(out)


def test_liste_sans_style_dans_un_document_neuf_sans_paragraphe_orphelin():
    """Le repli des listes (styles absents) ne laisse plus un paragraphe vide
    devant chaque élément : `add_paragraph(style=…)` crée le paragraphe
    avant d'échouer sur le style."""
    doc = Document()
    doc.styles["List Bullet"].delete()
    export._docx_riche(doc, "- un\n- deux")
    assert textes(doc) == ["• un", "• deux"]


def test_pdf_ignore_le_gabarit():
    cfg = cfg_avec(gabarit_porte_identite=True)
    cfg["mise_en_page"]["gabarit"] = {"nom": "x.docx"}
    pdf = export.to_pdf(BILAN, cfg)
    assert pdf.startswith(b"%PDF")


# --- routes -------------------------------------------------------------------

def _png() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (60, 30), (10, 20, 30)).save(buf, "PNG")
    return buf.getvalue()


def _cle_gabarit_presente() -> bool:
    with security.transaction() as con:
        row = con.execute(
            "SELECT length(value) FROM config WHERE key = ?", (config.ConfigStore.CLE_GABARIT,)
        ).fetchone()
    return bool(row and row[0])


def test_api_gabarit_depot_export_recuperation_retrait(client):
    fichiers = {"fichier": ("Papier en-tête.dotx", en_dotx(gabarit()), "application/octet-stream")}
    r = client.put("/api/config/gabarit", files=fichiers)
    assert r.status_code == 200, r.text
    g = r.json()["mise_en_page"]["gabarit"]
    assert g["nom"] == "Papier en-tete.docx" and g["en_tete"] is True and g["styles"][0] == "Normal"
    ov = client.get("/api/config/overrides").json()["mise_en_page"]
    assert ov["gabarit"]["nom"] == "Papier en-tete.docx" and "donnees" not in ov["gabarit"]
    assert _cle_gabarit_presente()
    # Récupération : le .docx converti, sous son nom.
    r = client.get("/api/config/gabarit")
    assert r.status_code == 200 and r.headers["content-type"].startswith("application/vnd.openxml")
    assert 'filename="Papier en-tete.docx"' in r.headers["content-disposition"]
    assert EN_TETE in membre(r.content, "word/header1.xml")
    # L'export Word d'un bilan part du gabarit ; le PDF ne change pas.
    bid = client.post("/api/bilans", json={"domaines": ["voix"]}).json()["id"]
    docx = client.get(f"/api/bilans/{bid}/export", params={"format": "docx"}).content
    assert EN_TETE in membre(docx, "word/header1.xml")
    assert Document(io.BytesIO(docx)).styles["Normal"].font.name == "Garamond"
    assert client.get(f"/api/bilans/{bid}/export", params={"format": "pdf"}).content.startswith(b"%PDF")
    # L'exemple Word de l'écran Paramètres aussi, avec les réglages envoyés.
    r = client.post("/api/config/mise_en_page/apercu", params={"format": "docx"},
                    json={"rubriques_numerotees": True, "gabarit_porte_identite": True})
    assert r.status_code == 200 and "exemple-mise-en-page.docx" in r.headers["content-disposition"]
    assert EN_TETE in membre(r.content, "word/header1.xml")
    assert "1. Anamnèse" in textes(Document(io.BytesIO(r.content)))
    assert client.post("/api/config/mise_en_page/apercu", json={}).headers["content-type"] == "application/pdf"
    # Le réglage « porte l'identité » passe par le PUT ordinaire, la
    # description du gabarit jamais.
    eff = client.put("/api/config", json={"overrides": {"mise_en_page": {"gabarit_porte_identite": True}}}).json()
    assert eff["mise_en_page"]["gabarit_porte_identite"] is True and eff["mise_en_page"]["gabarit"]["nom"]
    r = client.put("/api/config", json={"overrides": {"mise_en_page": {"gabarit": {"nom": "faux.docx"}}}})
    assert r.status_code == 422
    # Retour aux valeurs recommandées de la section, clé par clé : le gabarit reste.
    eff = client.delete("/api/config/mise_en_page", params={"cles": "gabarit_porte_identite,police"}).json()
    assert eff["mise_en_page"]["gabarit_porte_identite"] is False and eff["mise_en_page"]["gabarit"]["nom"]
    # Retrait : description et octets disparaissent, le reste des réglages tient.
    client.put("/api/config", json={"overrides": {"mise_en_page": {"police": "Arial"}}})
    eff = client.delete("/api/config/gabarit").json()
    assert eff["mise_en_page"]["gabarit"] is None and eff["mise_en_page"]["police"] == "Arial"
    assert "gabarit" not in client.get("/api/config/overrides").json()["mise_en_page"]
    assert not _cle_gabarit_presente()
    assert client.get("/api/config/gabarit").status_code == 404
    docx = client.get(f"/api/bilans/{bid}/export", params={"format": "docx"}).content
    assert EN_TETE not in membre(docx, "word/header1.xml")


def test_api_gabarit_refus_et_reinitialisation(client):
    r = client.put("/api/config/gabarit", files={"fichier": ("bilan.txt", b"du texte", "text/plain")})
    assert r.status_code == 422 and "document Word" in r.json()["detail"]
    trop = b"PK" + b"0" * export.TAILLE_MAX_GABARIT
    assert client.put("/api/config/gabarit", files={"fichier": ("g.docx", trop, "application/octet-stream")}).status_code == 413
    assert client.get("/api/config").json()["mise_en_page"]["gabarit"] is None
    # Un refus ne touche pas le gabarit en place.
    assert client.put("/api/config/gabarit", files={"fichier": ("g.docx", gabarit(), "application/octet-stream")}).status_code == 200
    client.put("/api/config/gabarit", files={"fichier": ("bilan.txt", b"du texte", "text/plain")})
    assert client.get("/api/config/gabarit").status_code == 200
    # La section effacée en bloc et la réinitialisation générale emportent
    # aussi les octets du gabarit : rien d'orphelin dans le coffre.
    client.delete("/api/config/mise_en_page")
    assert not _cle_gabarit_presente() and client.get("/api/config/gabarit").status_code == 404
    client.put("/api/config/gabarit", files={"fichier": ("g.docx", gabarit(), "application/octet-stream")})
    assert _cle_gabarit_presente()
    client.delete("/api/config")
    assert not _cle_gabarit_presente()
    assert client.get("/api/config").json()["mise_en_page"]["gabarit"] is None
