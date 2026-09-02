"""Tests de l'API HTTP (FastAPI TestClient) — hors ligne : LLM et embeddings mockés."""
from __future__ import annotations

import time

from app import config, llm, rag, security
from tests.conftest import PASSPHRASE

BILAN_TXT = (
    "Anamnèse\nEnfant né à terme, marche à 12 mois.\n\n"
    "Projet thérapeutique\nDeux séances par semaine."
)


# --- session / verrouillage ------------------------------------------------------

def test_statut_et_verrouillage(client):
    from app import __version__

    s = client.get("/api/status").json()
    assert s == {"db_exists": True, "unlocked": True, "first_run": False,
                 "version": __version__}
    assert s["version"].count(".") == 2
    assert client.post("/api/lock").status_code == 200
    assert client.get("/api/status").json()["unlocked"] is False
    # endpoint protégé -> 423
    assert client.get("/api/bilans").status_code == 423
    # mauvaise passphrase -> 401 ; bonne -> 200
    assert client.post("/api/unlock", json={"passphrase": "mauvaise"}).status_code == 401
    assert client.post("/api/unlock", json={"passphrase": PASSPHRASE}).status_code == 200
    assert client.post("/api/unlock", json={"passphrase": "  "}).status_code == 400


def test_auto_verrouillage_inactivite(client):
    client.put("/api/config", json={"overrides": {"rgpd": {"verrouillage_inactivite_minutes": 1}}})
    security._state["last_activity"] = time.monotonic() - 120
    assert client.get("/api/bilans").status_code == 423
    assert client.get("/api/status").json()["unlocked"] is False


def test_keepalive_rafraichit(client):
    assert client.post("/api/keepalive").status_code == 200


def test_verrou_entre_dependance_et_transaction(client, monkeypatch):
    """Coffre verrouillé APRÈS require_unlock mais AVANT transaction() (course
    multi-onglets via POST /api/lock) : 423 partout, plus jamais 500 (BUG-05)."""
    vrai_touch = security.touch

    def touch_puis_verrou():
        vrai_touch()
        security.lock()

    monkeypatch.setattr(security, "touch", touch_puis_verrou)
    r = client.get("/api/patients")
    assert r.status_code == 423
    assert "verrouillée" in r.json()["detail"]


def test_disque_plein_au_commit(client):
    """Un commit qui échoue (disque plein) doit donner un message actionnable,
    pas un 500 opaque (BUG-06)."""
    import sqlcipher3

    vrai_con = security._state["con"]

    class ConDisquePlein:
        def __getattr__(self, nom):
            return getattr(vrai_con, nom)

        def commit(self):
            # La classe de la base chiffrée réelle (≠ sqlite3.OperationalError).
            raise sqlcipher3.OperationalError("database or disk is full")

    security._state["con"] = ConDisquePlein()
    try:
        r = client.get("/api/patients")
    finally:
        security._state["con"] = vrai_con
    assert r.status_code == 503
    assert "espace disque" in r.json()["detail"]


def test_verrouillage_survit_config_corrompue(client):
    """Une vieille surcharge mal typée déjà stockée (avant la validation
    Pydantic) ne doit plus bloquer toutes les routes protégées (audit C5)."""
    with security.transaction() as con:
        config.ConfigStore(con).set_overrides(
            {"rgpd": {"verrouillage_inactivite_minutes": "quinze"}}
        )
    assert client.get("/api/bilans").status_code == 200


# --- config -----------------------------------------------------------------------

def test_config_get_put_delete(client):
    eff = client.get("/api/config").json()
    assert eff["llm"]["model"] == config.DEFAULTS["llm"]["model"]
    eff = client.put(
        "/api/config", json={"overrides": {"style": {"few_shot_k": 7}}}
    ).json()
    assert eff["style"]["few_shot_k"] == 7
    assert client.get("/api/config").json()["style"]["few_shot_k"] == 7
    eff = client.delete("/api/config").json()
    assert eff["style"]["few_shot_k"] == config.DEFAULTS["style"]["few_shot_k"]


def test_config_overrides_expose_les_surcharges_seules(client):
    assert client.get("/api/config/overrides").json() == {}
    client.put("/api/config", json={"overrides": {"llm": {"model": "x"}}})
    assert client.get("/api/config/overrides").json() == {"llm": {"model": "x"}}


def test_config_dictee_max_minutes(client):
    """La borne de durée de dictée (BUG-04) existe, se configure et est validée."""
    assert client.get("/api/config").json()["rgpd"]["dictee_max_minutes"] == 30
    eff = client.put(
        "/api/config", json={"overrides": {"rgpd": {"dictee_max_minutes": 10}}}
    ).json()
    assert eff["rgpd"]["dictee_max_minutes"] == 10
    assert client.put(
        "/api/config", json={"overrides": {"rgpd": {"dictee_max_minutes": -1}}}
    ).status_code == 422


def test_domaines_publics(client):
    doms = client.get("/api/domaines").json()
    assert {"cle": "voix", "titre": "Voix"} in doms


def test_trame_et_catalogue_configurables_via_api(client):
    client.put("/api/config", json={"overrides": {
        "trame": {"sections": [{"cle": "libre", "titre": "Rubrique libre"}]},
        "catalogues": {"voix": {"tests": [{"nom": "Échelle maison", "mesure": "m",
                                           "metriques": ["qualitatif"]}]}},
    }})
    b = client.post("/api/bilans", json={"domaines": ["voix"]}).json()
    assert [s["cle"] for s in b["sections"]] == ["libre"]
    cat = client.get("/api/catalogues/voix").json()
    assert [t["nom"] for t in cat["tests"]] == ["Échelle maison"]


# --- éditeurs dédiés : remplacement en bloc ---------------------------------------

def test_config_trame_remplacement_et_retour_defauts(client):
    """PUT /api/config/trame remplace EN BLOC (une rubrique retirée disparaît
    vraiment — impossible via la fusion du PUT /api/config) ; DELETE rend la
    trame réglementaire et ne laisse aucune surcharge figée."""
    deux = [{"cle": "libre", "titre": "Rubrique libre"},
            {"cle": "epreuves", "titre": "Épreuves"}]
    assert client.put("/api/config/trame", json={"sections": deux}).status_code == 200
    r = client.put("/api/config/trame", json={"sections": deux[:1]})
    assert r.status_code == 200
    ov = client.get("/api/config/overrides").json()
    assert ov["trame"]["sections"] == [{"cle": "libre", "titre": "Rubrique libre"}]
    b = client.post("/api/bilans", json={"domaines": []}).json()
    assert [s["cle"] for s in b["sections"]] == ["libre"]
    eff = client.delete("/api/config/trame").json()
    assert eff["trame"] == config.DEFAULTS["trame"]
    assert "trame" not in client.get("/api/config/overrides").json()
    b = client.post("/api/bilans", json={"domaines": []}).json()
    assert len(b["sections"]) == 7  # tronc commun réglementaire


def test_config_trame_validation(client):
    ko = [
        {"sections": []},                                     # liste vide
        {"sections": [{"cle": "x", "titre": ""}]},            # titre vide
        {"sections": [{"cle": "  ", "titre": "T"}]},          # clé blanche
        {"sections": [{"titre": "Sans clé"}]},                # clé absente
    ]
    for corps in ko:
        assert client.put("/api/config/trame", json=corps).status_code == 422
    # l'app reste utilisable après un rejet
    assert client.get("/api/patients").status_code == 200


def test_config_catalogues_remplacement_et_suppression_domaine(client):
    r = client.put("/api/config/catalogues", json={
        "voix": {"guidance": "Ma guidance.", "tests": [{"nom": "Échelle maison"}]},
    })
    assert r.status_code == 200
    cat = client.get("/api/catalogues/voix").json()
    assert cat["guidance"] == "Ma guidance."
    assert [t["nom"] for t in cat["tests"]] == ["Échelle maison"]
    # remplacement par {} : le domaine disparaît, le catalogue intégré revient
    assert client.put("/api/config/catalogues", json={}).status_code == 200
    assert "catalogues" not in client.get("/api/config/overrides").json()
    cat = client.get("/api/catalogues/voix").json()
    assert cat["guidance"] != "Ma guidance." and len(cat["tests"]) > 0
    # validations : nom manquant, métrique inconnue, domaine inconnu
    assert client.put("/api/config/catalogues", json={
        "voix": {"tests": [{"mesure": "sans nom"}]}}).status_code == 422
    assert client.put("/api/config/catalogues", json={
        "voix": {"tests": [{"nom": "T", "metriques": ["score_brut"]}]}}).status_code == 422
    r = client.put("/api/config/catalogues", json={"telepathie": {"guidance": "x"}})
    assert r.status_code == 422 and "Domaine inconnu" in r.json()["detail"]


def test_config_prompts_remplacement_et_vide_efface(client):
    assert client.put(
        "/api/config/prompts", json={"structure_system": "MA CONSIGNE {cles}"}
    ).status_code == 200
    assert client.get("/api/config/overrides").json()["prompts"] == {
        "structure_system": "MA CONSIGNE {cles}"
    }
    # vide = retour à la consigne intégrée, pas de surcharge vide figée
    assert client.put("/api/config/prompts", json={"structure_system": "  "}).status_code == 200
    assert "prompts" not in client.get("/api/config/overrides").json()
    client.put("/api/config/prompts", json={"structure_system": "X"})
    client.delete("/api/config/prompts")
    assert "prompts" not in client.get("/api/config/overrides").json()
    # les remplacements sont audités
    with security.transaction() as con:
        actions = [r[0] for r in con.execute("SELECT action FROM audit_log").fetchall()]
    assert "config_prompts" in actions


def test_prompt_defaut_expose_et_reutilisable(client, monkeypatch, mock_embed):
    """Verrou anti-régression du piège des accolades : la consigne intégrée
    exposée (accolades simples) doit, enregistrée telle quelle, produire
    EXACTEMENT le même prompt système que le défaut (.format)."""
    from app import prompts as prompts_mod

    texte = client.get("/api/prompts/structure-defaut").json()["prompt"]
    assert "{cles}" in texte and '{"updates"' in texte
    assert "{{" not in texte and "}}" not in texte
    client.put("/api/config/prompts", json={"structure_system": texte})

    captures = {}

    async def chat_espion(system, user, **kw):
        captures["system"] = system
        return '{"updates": [], "questions": []}'

    monkeypatch.setattr(llm, "chat_json", chat_espion)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    assert client.post(
        f"/api/bilans/{bid}/structure", json={"transcription": "Texte."}
    ).status_code == 200
    cles = ", ".join(
        sorted(s["cle"] for s in client.get(f"/api/bilans/{bid}").json()["sections"])
    )
    assert captures["system"] == prompts_mod.STRUCTURE_SYSTEM.format(cles=cles)


def test_statut_valide_puis_envoye(client):
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    b = client.put(f"/api/bilans/{bid}/statut", json={"statut": "valide"}).json()
    assert b["statut"] == "valide"
    b = client.put(
        f"/api/bilans/{bid}/statut",
        json={"statut": "envoye", "destinataire": "Dr Martin"},
    ).json()
    assert b["statut"] == "envoye"
    assert client.get("/api/bilans").json()[0]["statut"] == "envoye"
    assert client.put("/api/bilans/999/statut", json={"statut": "valide"}).status_code == 404
    # statut hors énumération -> 422 (validation Pydantic)
    assert client.put(f"/api/bilans/{bid}/statut", json={"statut": "nimporte"}).status_code == 422


# --- bilans -----------------------------------------------------------------------

def test_bilan_patient_inexistant(client):
    """Créer un bilan pour un patient inconnu → 404 clair, pas un 500 sur la
    contrainte de clé étrangère (BUG-01)."""
    r = client.post("/api/bilans", json={"domaines": [], "patient_id": 99999})
    assert r.status_code == 404
    assert "Patient introuvable" in r.json()["detail"]


def test_parcours_bilan_complet(client):
    b = client.post(
        "/api/bilans", json={"domaines": ["langage_ecrit"], "type": "initial_complexe"}
    ).json()
    bid = b["id"]
    assert [s["cle"] for s in b["sections"]][0] == "administratif"
    assert client.get("/api/bilans").json()[0]["id"] == bid
    assert client.get("/api/bilans/999").status_code == 404

    # édition + validation d'une rubrique
    r = client.put(
        f"/api/bilans/{bid}/sections/anamnese",
        json={"contenu": "Texte relu.", "statut": "valide"},
    )
    assert r.status_code == 200
    assert client.put(
        f"/api/bilans/{bid}/sections/inconnue", json={"contenu": "x"}
    ).status_code == 404

    # épreuve avec drapeau automatique
    b = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "Alouette-R", "domaine": "langage_ecrit",
        "resultats": [{"score_brut": "112", "etalonnage_type": "percentile",
                       "etalonnage_valeur": "5"}],
    }).json()
    assert b["epreuves"][0]["resultats"][0]["drapeau_seuil"] == "pathologique"

    # cotation (type complexe -> AMO 34)
    cot = client.post(f"/api/bilans/{bid}/cotation").json()
    assert cot["code_amo"] == "AMO 34" and cot["montant"] == round(34 * 2.60, 2)

    # exports
    md = client.get(f"/api/bilans/{bid}/export?format=md")
    assert "## Anamnèse" in md.text and "Texte relu." in md.text
    docx = client.get(f"/api/bilans/{bid}/export?format=docx")
    assert docx.content[:2] == b"PK"  # zip Office valide
    assert "ANAMNÈSE" in client.get(f"/api/bilans/{bid}/export?format=txt").text


# --- audit 2026-08-11, lot 1 : rien de faux ni d'indélébile dans le document --

def test_epreuve_sans_resultat_exploitable_refusee(client):
    """Un corps sans résultat utile créait une épreuve coquille (200), soit une
    ligne vide dans le tableau du compte-rendu."""
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    assert client.post(f"/api/bilans/{bid}/epreuves",
                       json={"test_nom": "Alouette-R"}).status_code == 422
    assert client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "Alouette-R", "resultats": [{"sous_epreuve": "lecture"}],
    }).status_code == 422
    # Un résultat purement qualitatif (interprétation seule) reste accepté.
    r = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "GRBAS", "resultats": [{"interpretation": "voix soufflée"}],
    })
    assert r.status_code == 200 and len(r.json()["epreuves"]) == 1


def test_epreuve_supprimable(client):
    """Une échelle mal choisie produit un drapeau faux : il doit pouvoir partir
    autrement qu'en supprimant le patient entier."""
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    b = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "Alouette-R",
        "resultats": [{"score_brut": "112", "etalonnage_type": "percentile",
                       "etalonnage_valeur": "5"}],
    }).json()
    eid = b["epreuves"][0]["id"]
    autre = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    # Une épreuve ne se supprime que depuis le bilan qui la porte.
    assert client.delete(f"/api/bilans/{autre}/epreuves/{eid}").status_code == 404
    r = client.delete(f"/api/bilans/{bid}/epreuves/{eid}")
    assert r.status_code == 200 and r.json()["epreuves"] == []
    assert client.delete(f"/api/bilans/{bid}/epreuves/{eid}").status_code == 404
    # Le résultat suit l'épreuve (cascade), le tableau d'export est vide.
    assert "Résultats des épreuves" not in client.get(
        f"/api/bilans/{bid}/export?format=md").text


def test_bilan_supprimable_sans_toucher_au_patient(client):
    p = client.post("/api/patients", json={"nom": "Durand"}).json()
    bid = client.post("/api/bilans",
                      json={"domaines": [], "patient_id": p["id"],
                            "prescripteur": "Bernard"}).json()["id"]
    assert client.delete(f"/api/bilans/{bid}").status_code == 200
    assert client.get(f"/api/bilans/{bid}").status_code == 404
    assert client.delete(f"/api/bilans/{bid}").status_code == 404
    # Le patient reste : supprimer un bilan n'est pas un effacement RGPD.
    assert any(x["id"] == p["id"] for x in client.get("/api/patients").json())


def test_etalonnage_hors_bornes_signale_sans_bloquer(client):
    """On signale, on ne corrige pas : le drapeau reste calculé, mais la saisie
    invraisemblable est nommée (percentile -300, note 85 sur l'échelle moy. 10)."""
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "Alouette-R",
        "resultats": [{"score_brut": "12", "etalonnage_type": "percentile",
                       "etalonnage_valeur": "-300"}],
    })
    assert r.status_code == 200
    avert = r.json()["avertissements"]
    assert len(avert) == 1 and "0 à 100" in avert[0]

    r = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "EXALANG", "resultats": [
            {"sous_epreuve": "vocabulaire", "score_brut": "85",
             "etalonnage_type": "note_standard", "etalonnage_valeur": "85"}],
    })
    assert "1 à 19" in r.json()["avertissements"][0]

    # Saisie plausible : aucun bruit.
    r = client.post(f"/api/bilans/{bid}/epreuves", json={
        "test_nom": "Vineland", "resultats": [
            {"etalonnage_type": "note_standard_100", "etalonnage_valeur": "85"}],
    })
    assert r.json()["avertissements"] == []


def test_seuils_incoherents_refuses(client):
    """Saisir 1.5 au lieu de -1.5 basculait tous les résultats normaux en
    « zone de fragilité », sans un mot."""
    seuils = {"fragilite_et": 1.5, "pathologique_et": -1.5, "severe_et": -2}
    r = client.put("/api/config", json={"overrides": {"seuils": seuils}})
    assert r.status_code == 422
    # Ordre incohérent : sévère au-dessus de pathologique.
    r = client.put("/api/config", json={"overrides": {"seuils": {
        "fragilite_et": -1, "pathologique_et": -2, "severe_et": -1.5}}})
    assert r.status_code == 422
    # Percentiles : même contrôle, dans l'autre sens.
    r = client.put("/api/config", json={"overrides": {"seuils": {
        "fragilite_percentile": 16, "pathologique_percentile": 2,
        "severe_percentile": 7}}})
    assert r.status_code == 422
    # Jeu cohérent : accepté.
    r = client.put("/api/config", json={"overrides": {"seuils": {
        "fragilite_et": -1, "pathologique_et": -1.5, "severe_et": -2}}})
    assert r.status_code == 200


def test_trame_cles_dupliquees_refusees(client):
    """Deux rubriques de même clé : le document imprimait la rubrique deux fois
    et le texte de l'IA n'atterrissait que dans l'une d'elles."""
    r = client.put("/api/config/trame", json={"sections": [
        {"cle": "anamnese", "titre": "Anamnèse"},
        {"cle": "anamnese", "titre": "Anamnèse (suite)"},
    ]})
    assert r.status_code == 422


def test_export_format_inconnu_refuse(client):
    """« ?format=exe » renvoyait du Markdown en 200."""
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    assert client.get(f"/api/bilans/{bid}/export?format=exe").status_code == 422
    assert client.get(f"/api/bilans/{bid}/export?format=pdf").status_code == 200


def test_export_brouillon_porte_la_mention(client):
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    client.put(f"/api/bilans/{bid}/sections/anamnese", json={"contenu": "Texte."})
    assert "BROUILLON" in client.get(f"/api/bilans/{bid}/export?format=md").text
    client.put(f"/api/bilans/{bid}/statut", json={"statut": "valide"})
    assert "BROUILLON" not in client.get(f"/api/bilans/{bid}/export?format=md").text


def test_catalogue_par_domaine(client):
    cat = client.get("/api/catalogues/langage_ecrit").json()
    assert any(t["nom"] == "Alouette-R" for t in cat["tests"])
    # domaine inconnu -> guidance générique, pas d'erreur
    assert client.get("/api/catalogues/inconnu").json()["tests"] == []


# --- références (import + RAG) ------------------------------------------------------

def test_references_import_liste_suppression(client, mock_embed):
    r = client.post(
        "/api/references",
        files={"file": ("bilan.txt", BILAN_TXT.encode(), "text/plain")},
        data={"domaine": "langage_oral"},
    )
    assert r.status_code == 200 and r.json()["n"] == 2
    refs = client.get("/api/references").json()
    assert len(refs) == 2 and {x["section_cle"] for x in refs} == {"anamnese", "projet"}
    assert client.delete(f"/api/references/{refs[0]['id']}").status_code == 200
    assert len(client.get("/api/references").json()) == 1
    # fichier sans texte -> 400
    r = client.post("/api/references", files={"file": ("v.txt", b"  ", "text/plain")})
    assert r.status_code == 400


def test_import_docx(client, mock_embed):
    """Le .docx — format que l'app exporte elle-même — doit s'importer en
    texte lisible, pas en binaire ZIP vectorisé (audit)."""
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("Anamnèse")
    doc.add_paragraph("Enfant né à terme, marche à 12 mois.")
    buf = io.BytesIO()
    doc.save(buf)
    r = client.post(
        "/api/references",
        files={"file": ("bilan.docx", buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"domaine": "langage_oral"},
    )
    assert r.status_code == 200 and r.json()["n"] >= 1
    refs = client.get("/api/references").json()
    assert "anamnese" in {x["section_cle"] for x in refs}


def test_import_odt(client, mock_embed):
    """Le .odt — format par défaut de LibreOffice, courant chez les
    praticiens — doit s'importer comme le .docx."""
    import io
    import zipfile

    content = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<office:document-content'
        ' xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">'
        "<office:body><office:text>"
        "<text:h>Anamnèse</text:h>"
        "<text:p>Enfant né à terme, marche à 12 mois.</text:p>"
        "</office:text></office:body></office:document-content>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        z.writestr("content.xml", content)
    r = client.post(
        "/api/references",
        files={"file": ("bilan.odt", buf.getvalue(),
                        "application/vnd.oasis.opendocument.text")},
        data={"domaine": "langage_oral"},
    )
    assert r.status_code == 200 and r.json()["n"] >= 1
    refs = client.get("/api/references").json()
    assert "anamnese" in {x["section_cle"] for x in refs}
    # .odt corrompu -> 400 explicite, comme le PDF corrompu
    r = client.post("/api/references",
                    files={"file": ("casse.odt", b"PK\x03\x04casse", "application/octet-stream")})
    assert r.status_code == 400 and ".odt illisible" in r.json()["detail"]


def test_export_docx_reimportable(client, mock_embed):
    """Aller-retour complet : un bilan exporté en Word se réimporte tel quel."""
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    client.put(f"/api/bilans/{bid}/sections/anamnese",
               json={"contenu": "Enfant né à terme."})
    data = client.get(f"/api/bilans/{bid}/export?format=docx").content
    r = client.post("/api/references", files={"file": ("bilan.docx", data, "application/octet-stream")})
    assert r.status_code == 200 and r.json()["n"] >= 1


def test_import_reference_pseudonymise_et_ecarte_l_entete(client, mock_embed):
    """Audit 2026-08-11 (3.1) : ces extraits sont relus par le modèle pendant
    la rédaction du bilan d'un AUTRE patient. Le bloc d'identité en tête de
    document (extrait « global », exempté du filtre de rubrique) ne doit pas y
    entrer, et ce qui reste est caviardé."""
    doc = (
        "COMPTE RENDU\n"
        "Patient : DURAND Léa, née le 12/03/2018\n"
        "Adresse : 12 rue des Lilas, 44000 Nantes\n"
        "Adressé par le Dr Bernard\n\n"
        "Anamnèse\n"
        "Léa est en CE1. Sa mère rapporte des difficultés depuis la GS.\n\n"
        "Projet thérapeutique\n"
        "Deux séances hebdomadaires."
    )
    r = client.post("/api/references",
                    files={"file": ("bilan.txt", doc.encode(), "text/plain")})
    assert r.status_code == 200
    j = r.json()
    assert j["extraits_ecartes"] == 1 and "global" not in j["sections"]
    assert j["elements_caviardes"] >= 1
    textes = " ".join(x["titre"] + x["section_cle"] for x in client.get("/api/references").json())
    assert "DURAND" not in textes
    # Le journal d'audit ne porte pas le nom du fichier (il n'est nettoyé par
    # aucune suppression).
    with security.transaction() as con:
        details = " ".join(
            r[0] or "" for r in con.execute(
                "SELECT details FROM audit_log WHERE action='import_reference'")
        )
    assert "bilan.txt" not in details and ".txt" in details


def test_import_reference_rattache_a_un_patient(client, mock_embed):
    p = client.post("/api/patients", json={"nom": "Durand"}).json()
    r = client.post("/api/references",
                    files={"file": ("ref.txt", b"Anamnese\nUn texte de style.", "text/plain")},
                    data={"patient_id": str(p["id"])})
    assert r.status_code == 200 and r.json()["n"] >= 1
    assert client.get("/api/patients").json()[0]["nb_references"] >= 1
    # Effacement RGPD : les extraits rattachés partent avec le dossier.
    assert client.delete(f"/api/patients/{p['id']}").status_code == 200
    assert client.get("/api/references").json() == []
    # Patient inexistant : refus explicite plutôt qu'un rattachement fantôme.
    assert client.post("/api/references",
                       files={"file": ("ref.txt", b"Texte.", "text/plain")},
                       data={"patient_id": "999"}).status_code == 404


def test_import_binaire_rejete(client, mock_embed):
    # extension inconnue -> refus explicite
    r = client.post("/api/references",
                    files={"file": ("archive.zip", b"PK\x03\x04xxxx", "application/zip")})
    assert r.status_code == 400 and "pris en charge" in r.json()["detail"]
    # binaire déguisé en .txt -> refus (octet nul)
    r = client.post("/api/references",
                    files={"file": ("piege.txt", b"abc\x00def", "text/plain")})
    assert r.status_code == 400


def test_pack_exemples_import_remplacement_retrait(client, mock_embed):
    """Le pack embarqué s'indexe en un clic, se remplace sans doublon au
    re-clic, et se retire sans toucher aux bilans importés par le praticien."""
    # Une référence du praticien, qui doit survivre à toutes les opérations.
    r = client.post(
        "/api/references",
        files={"file": ("mien.txt", BILAN_TXT.encode(), "text/plain")},
        data={"domaine": "langage_oral"},
    )
    n_perso = r.json()["n"]

    r = client.post("/api/references/pack")
    assert r.status_code == 200
    s = r.json()
    assert s["n_fichiers"] == 11
    assert s["n_extraits"] >= s["n_fichiers"] * 5  # ≥ 5 rubriques par bilan
    refs = client.get("/api/references").json()
    assert len([x for x in refs if x["source"] == "fictif"]) == s["n_extraits"]

    # Re-clic : remplacement à l'identique, jamais d'addition.
    s2 = client.post("/api/references/pack").json()
    refs2 = client.get("/api/references").json()
    assert s2["n_extraits"] == s["n_extraits"]
    assert len([x for x in refs2 if x["source"] == "fictif"]) == s["n_extraits"]
    assert len([x for x in refs2 if x["source"] == "import"]) == n_perso

    # Retrait du pack : seuls les extraits fictifs disparaissent.
    assert client.delete("/api/references/pack").json()["n"] == s["n_extraits"]
    refs3 = client.get("/api/references").json()
    assert len(refs3) == n_perso and all(x["source"] == "import" for x in refs3)
    # Retrait d'un pack déjà retiré : 0, sans erreur.
    assert client.delete("/api/references/pack").json()["n"] == 0


def test_import_pdf_corrompu(client, mock_embed):
    """PDF corrompu ou protégé : 400 explicite au lieu d'un 500 pypdf (BUG-03)."""
    r = client.post(
        "/api/references",
        files={"file": ("bilan.pdf", b"%PDF-1.4 corrompu", "application/pdf")},
    )
    assert r.status_code == 400
    assert "PDF illisible" in r.json()["detail"]


def test_bilans_pagination(client):
    ids = [client.post("/api/bilans", json={"domaines": []}).json()["id"] for _ in range(3)]
    page1 = client.get("/api/bilans?limit=2").json()
    assert [b["id"] for b in page1] == [ids[2], ids[1]]
    page2 = client.get("/api/bilans?limit=2&offset=2").json()
    assert [b["id"] for b in page2] == [ids[0]]
    # bornes défensives
    assert client.get("/api/bilans?limit=0").status_code == 200
    assert client.get("/api/bilans?offset=-1").status_code == 200


def test_references_embeddings_indisponibles(client, monkeypatch):
    def boom(text, cfg):
        raise rag.EmbeddingUnavailable("modèle absent")

    monkeypatch.setattr(rag, "embed", boom)
    r = client.post(
        "/api/references", files={"file": ("b.txt", BILAN_TXT.encode(), "text/plain")}
    )
    assert r.status_code == 503


# --- patients --------------------------------------------------------------------

def test_patients_api_et_cascade(client):
    assert client.post("/api/patients", json={"nom": "  "}).status_code == 400
    p = client.post("/api/patients", json={
        "nom": "Durand", "prenom": "Léa", "date_naissance": "2018-03-12", "sexe": "F",
    }).json()
    bid = client.post("/api/bilans", json={"domaines": [], "patient_id": p["id"]}).json()["id"]
    # le bilan expose l'identité ; la liste des patients compte les bilans
    assert client.get(f"/api/bilans/{bid}").json()["patient"]["nom"] == "Durand"
    assert client.get("/api/bilans").json()[0]["patient_nom"] == "Durand"
    assert client.get("/api/patients").json()[0]["nb_bilans"] == 1
    # mise à jour
    r = client.put(f"/api/patients/{p['id']}", json={"nom": "Durand", "prenom": "Léa-Marie"})
    assert r.json()["prenom"] == "Léa-Marie"
    assert client.put("/api/patients/999", json={"nom": "X"}).status_code == 404
    # effacement RGPD : le bilan rattaché disparaît
    assert client.delete(f"/api/patients/{p['id']}").status_code == 200
    assert client.get(f"/api/bilans/{bid}").status_code == 404
    assert client.delete("/api/patients/999").status_code == 404


def test_export_contient_le_patient(client):
    p = client.post("/api/patients", json={
        "nom": "Durand", "prenom": "Léa", "date_naissance": "2018-03-12",
    }).json()
    bid = client.post("/api/bilans", json={"domaines": [], "patient_id": p["id"]}).json()["id"]
    md = client.get(f"/api/bilans/{bid}/export?format=md").text
    # sexe non renseigné : pas de participe accordé, pas de « né(e) » non plus
    assert "Patient : DURAND Léa, date de naissance : 12/03/2018" in md
    # sexe renseigné : accord depuis le dossier
    client.put(f"/api/patients/{p['id']}", json={
        "nom": "Durand", "prenom": "Léa", "date_naissance": "2018-03-12", "sexe": "F",
    })
    md = client.get(f"/api/bilans/{bid}/export?format=md").text
    assert "Patient : DURAND Léa, née le 12/03/2018" in md


# --- sauvegarde chiffrée ------------------------------------------------------------

def test_sauvegarde_api(client):
    r = client.post("/api/sauvegarde")
    assert r.status_code == 200
    s = r.json()
    assert s["octets"] > 0 and "bilan-ortho-sauvegarde-" in s["fichier"]
    from pathlib import Path

    etat = client.get("/api/sauvegardes").json()
    assert etat["derniere"] is not None
    assert any(f["fichier"] == Path(s["fichier"]).name for f in etat["fichiers"])


def test_restauration_api(client):
    """Parcours complet de restauration via l'API + demandes invalides."""
    from pathlib import Path

    client.put("/api/config", json={"overrides": {"sauvegarde": {"auto_jours": 0}}})
    nom = Path(client.post("/api/sauvegarde").json()["fichier"]).name
    r = client.put("/api/config", json={"overrides": {"llm": {"model": "apres-sauvegarde"}}})
    assert r.status_code == 200
    r = client.post("/api/restauration", json={"fichier": nom, "passphrase": PASSPHRASE})
    assert r.status_code == 200
    corps = r.json()
    assert corps["ok"] is True and corps["filet"].startswith("bilan-ortho-sauvegarde-")
    # la surcharge postérieure à la sauvegarde a disparu : base bien remplacée,
    # et l'app est restée utilisable sans re-déverrouillage manuel
    assert client.get("/api/config").json()["llm"]["model"] != "apres-sauvegarde"
    # passphrase incorrecte → 400 explicite, app toujours utilisable
    r = client.post("/api/restauration", json={"fichier": nom, "passphrase": "mauvaise"})
    assert r.status_code == 400 and "passphrase" in r.json()["detail"]
    # nom hostile ou fichier inconnu → 400
    for nom_ko in ["../bilan.db", "bilan-ortho-sauvegarde-inexistante.db"]:
        r = client.post("/api/restauration", json={"fichier": nom_ko, "passphrase": PASSPHRASE})
        assert r.status_code == 400
    # passphrase vide → 400
    r = client.post("/api/restauration", json={"fichier": nom, "passphrase": " "})
    assert r.status_code == 400
    assert client.get("/api/status").json()["unlocked"] is True


def test_restauration_concurrente_409(client):
    """Une restauration déjà en cours → 409 explicite, pas de gel silencieux
    sur le verrou global."""
    from app import main as main_mod

    assert main_mod._restauration_verrou.acquire(blocking=False)
    try:
        r = client.post(
            "/api/restauration",
            json={"fichier": "bilan-ortho-sauvegarde-x.db", "passphrase": PASSPHRASE},
        )
    finally:
        main_mod._restauration_verrou.release()
    assert r.status_code == 409
    assert "déjà en cours" in r.json()["detail"]


# --- structuration (LLM mocké) -------------------------------------------------------

def test_structure_avec_llm_mocke(client, monkeypatch, mock_embed):
    reponse = (
        '{"updates":[{"section":"anamnese","texte":"Né à terme."},'
        '{"section":"hors_trame","texte":"écarté"}],'
        '"questions":[{"section":"anamnese","question":"Quel âge ?","pourquoi":"étalonnage"}]}'
    )
    captured = {}

    async def fake_chat_json(system, user, **kw):
        captured["system"], captured["user"] = system, user
        return reponse

    monkeypatch.setattr(llm, "chat_json", fake_chat_json)

    # une référence importée au préalable doit nourrir le style
    client.post(
        "/api/references",
        files={"file": ("ref.txt", "Nous recevons le jeune L., très volontaire.".encode(), "text/plain")},
        data={"domaine": "langage_oral"},
    )
    bid = client.post("/api/bilans", json={"domaines": ["langage_oral"]}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Il est né à terme."})
    assert r.status_code == 200
    res = r.json()
    sections = {s["cle"]: s for s in res["bilan"]["sections"]}
    assert sections["anamnese"]["contenu"] == "Né à terme."
    assert sections["anamnese"]["statut"] == "propose_ia"
    assert res["questions"] == [
        {"section": "anamnese", "question": "Quel âge ?", "pourquoi": "étalonnage"}
    ]
    # la clé hors trame a été filtrée
    assert "hors_trame" not in sections
    # le prompt contient bien l'extrait de style et les préférences
    assert "Nous recevons le jeune L." in captured["user"]
    assert "vouvoyant" in captured["user"]
    # transcription vide -> 400
    assert client.post(f"/api/bilans/{bid}/structure", json={"transcription": " "}).status_code == 400

    # prompt de structuration personnalisé : utilisé tel quel, {cles} substitué
    client.put("/api/config", json={"overrides": {
        "prompts": {"structure_system": "MON PROMPT. Clés : {cles}."}
    }})
    client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Suite."})
    assert captured["system"].startswith("MON PROMPT.")
    assert "anamnese" in captured["system"]

    # avec un patient rattaché : l'âge (jamais le nom) est fourni au LLM
    p = client.post("/api/patients", json={
        "nom": "Durand", "prenom": "Léa", "date_naissance": "2018-03-12", "sexe": "F",
    }).json()
    bid2 = client.post(
        "/api/bilans", json={"domaines": ["langage_oral"], "patient_id": p["id"]}
    ).json()["id"]
    client.post(f"/api/bilans/{bid2}/structure", json={"transcription": "Elle est née à terme."})
    assert "âge à la date du bilan" in captured["user"]
    assert "sexe : F" in captured["user"] and "Ne pose PAS" in captured["user"]
    assert "Durand" not in captured["user"] and "Léa" not in captured["user"]


def test_structure_reponses_sans_dictee(client, monkeypatch, mock_embed):
    captured = {}

    async def fake_chat_json(system, user, **kw):
        captured["user"], captured["kw"] = user, kw
        return ('{"updates":[{"section":"anamnese",'
                '"texte":"Le patient est âgé de 7 ans."}],"questions":[]}')

    monkeypatch.setattr(llm, "chat_json", fake_chat_json)
    bid = client.post("/api/bilans", json={"domaines": ["langage_oral"]}).json()["id"]

    # ni dictée ni réponse -> 400
    assert client.post(
        f"/api/bilans/{bid}/structure", json={"transcription": " "}
    ).status_code == 400

    r = client.post(f"/api/bilans/{bid}/structure", json={
        "transcription": "",
        "reponses": [
            {"question": "Quel âge a le patient ?", "reponse": "7 ans", "section": "anamnese"},
        ],
        "questions_en_attente": ["Le score ELO est-il en note standard ?"],
        "questions_ecartees": ["Y a-t-il un suivi ORL ?"],
        "questions_repondues": ["Des antécédents familiaux ?"],
    })
    assert r.status_code == 200
    u = captured["user"]
    # la réponse et sa question arrivent structurées, avec la rubrique visée
    assert "Quel âge a le patient ?" in u and "7 ans" in u
    assert "rubrique visée : anamnese" in u
    # la mémoire du dialogue est transmise au LLM
    assert "EN ATTENTE" in u and "note standard" in u
    assert "ÉCARTÉES" in u and "suivi ORL" in u
    assert "DÉJÀ RÉPONDUES" in u and "antécédents familiaux" in u
    # pas de dictée ce tour-ci -> pas de bloc transcription
    assert "Transcription de la dictée" not in u
    # num_ctx par défaut transmis à Ollama (le prompt embarque tout le bilan)
    assert captured["kw"].get("num_ctx") == 8192
    # timeout borné transmis (un Ollama gelé ne suspend plus l'UI à l'infini)
    assert captured["kw"].get("timeout_s") == 600
    # la réponse est intégrée à la rubrique
    sections = {s["cle"]: s for s in r.json()["bilan"]["sections"]}
    assert sections["anamnese"]["contenu"] == "Le patient est âgé de 7 ans."

    # au tour suivant, le contenu déjà rédigé est visible dans le prompt
    client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Suite."})
    assert "« Le patient est âgé de 7 ans. »" in captured["user"]


def test_structure_modele_absent(client, monkeypatch, mock_embed):
    """Modèle non téléchargé : message qui pointe vers Paramètres → Modèles,
    pas un « Ollama injoignable » faux (BUG-02, UX-02)."""
    import httpx as _httpx

    async def chat_404(*a, **k):
        raise llm.ModeleIntrouvable("fantome:1b")

    monkeypatch.setattr(llm, "chat_json", chat_404)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 503
    assert "fantome:1b" in r.json()["detail"]
    assert "n'est pas téléchargé" in r.json()["detail"]

    async def chat_refuse(*a, **k):
        raise _httpx.ConnectError("connexion refusée")

    monkeypatch.setattr(llm, "chat_json", chat_refuse)
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 503
    assert "ne répond pas" in r.json()["detail"]


def test_structure_reponse_illisible(client, monkeypatch, mock_embed):
    """Réponse LLM sans JSON lisible : erreur explicite, pas un succès vide
    indistinguable d'un « rien à ajouter » (BUG-11)."""
    async def chat_blabla(*a, **k):
        return "blabla sans json"

    monkeypatch.setattr(llm, "chat_json", chat_blabla)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 502
    assert "n'a pas pu être lue" in r.json()["detail"]

    # JSON valide avec listes vides : succès légitime (0 mise à jour)
    async def chat_vide(*a, **k):
        return '{"updates": [], "questions": []}'

    monkeypatch.setattr(llm, "chat_json", chat_vide)
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 200
    assert r.json()["questions"] == []


def test_structure_analyse_deja_en_cours(client, mock_embed):
    """Une analyse déjà en cours sur le même bilan → 409 explicite (BUG-09)."""
    from app import main as main_mod

    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    main_mod._analyses_en_cours.add(bid)
    try:
        r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    finally:
        main_mod._analyses_en_cours.discard(bid)
    assert r.status_code == 409
    assert "déjà en cours" in r.json()["detail"]


def test_structure_concurrente_appliquee_une_fois(client, monkeypatch, mock_embed):
    """Deux analyses simultanées sur le même bilan : une seule passe (l'autre
    reçoit 409) et le contenu n'est appliqué qu'une fois — apply_updates fait
    un append, un doublon dupliquerait le texte (BUG-09)."""
    import asyncio
    import threading as th

    async def structure_lente(*a, **k):
        await asyncio.sleep(0.3)
        return {"updates": [{"section": "anamnese", "texte": "Ajout unique."}],
                "questions": []}

    monkeypatch.setattr(llm, "structure", structure_lente)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    codes = []

    def poster():
        codes.append(client.post(
            f"/api/bilans/{bid}/structure", json={"transcription": "Texte."}
        ).status_code)

    threads = [th.Thread(target=poster) for _ in range(2)]
    for t in threads:
        t.start()
    for t in threads:
        t.join()
    assert sorted(codes) == [200, 409]
    contenu = next(
        s["contenu"] for s in client.get(f"/api/bilans/{bid}").json()["sections"]
        if s["cle"] == "anamnese"
    )
    assert contenu.count("Ajout unique.") == 1


def test_structure_signale_rubriques_tronquees(client, monkeypatch, mock_embed):
    """Rubrique plus longue que llm.max_car_section : transmise partiellement
    au modèle ET signalée dans la réponse ; rien à signaler sinon (BUG-14)."""
    captured = {}

    async def fake_chat(system, user, **kw):
        captured["user"] = user
        return '{"updates": [], "questions": []}'

    monkeypatch.setattr(llm, "chat_json", fake_chat)
    client.put("/api/config", json={"overrides": {"llm": {"max_car_section": 120}}})
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Court."})
    assert r.status_code == 200 and r.json()["rubriques_tronquees"] == []
    client.put(f"/api/bilans/{bid}/sections/anamnese", json={"contenu": "long " * 100})
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.json()["rubriques_tronquees"] == ["anamnese"]
    # la troncature est bien effective dans le prompt (début conservé + […])
    assert "[…]" in captured["user"]


def test_structure_signale_test_substitue_et_prose_inventee(client, monkeypatch, mock_embed):
    """Les deux angles morts du garde-fou, mesurés le 2026-08-11 : un nom de
    test pris dans le catalogue du prompt, et une rubrique qui ne doit rien à
    la dictée. Aucun des deux ne portait de chiffre exploitable."""
    reponse = (
        '{"updates":[{"section":"epreuves","texte":"Alouette-R et EVALEO 6-15 '
        '(dictée de la Batelem au percentile cinq)."},'
        '{"section":"anamnese","texte":"Plainte actuelle : le patient rapporte des '
        "difficultés à la passation d'activités langagières, notamment lorsqu'il doit "
        'parler librement ou s\'exprimer devant un groupe. Il évite ces situations et '
        'cela lui cause une certaine souffrance."}],"questions":[]}'
    )

    async def fake_chat(system, user, **kw):
        return reponse

    monkeypatch.setattr(llm, "chat_json", fake_chat)
    bid = client.post("/api/bilans", json={"domaines": ["langage_ecrit"]}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={
        "transcription": "J'ai fait l'Alouette, elle lit 112 mots. En orthographe, "
                         "dictée de la Batelem, elle est au percentile cinq.",
    })
    assert r.status_code == 200
    par_section = {c["section"]: " ".join(c["signalements"])
                   for c in r.json()["rubriques_a_verifier"]}
    # Le test substitué est nommé — et « Alouette-R », dicté sans son suffixe,
    # ne l'est pas.
    assert "EVALEO 6-15" in par_section["epreuves"]
    assert "Alouette" not in par_section["epreuves"]
    # La rubrique inventée est signalée comme telle, sans qu'aucun chiffre
    # n'ait pu la trahir.
    assert "très peu adossée" in par_section["anamnese"]


def test_structure_signale_le_style_indisponible(client, monkeypatch, mock_embed):
    """Embeddings en panne : l'analyse aboutit quand même, mais la perte du
    style du praticien est signalée au lieu d'être avalée en silence. Rien
    n'est signalé à qui n'a importé aucun bilan de référence (pas de bruit)."""
    async def fake_chat(system, user, **kw):
        return '{"updates": [], "questions": []}'

    monkeypatch.setattr(llm, "chat_json", fake_chat)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]

    # aucune référence importée : pas de style attendu, donc rien à signaler
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 200 and r.json()["style_indisponible"] == ""

    client.post(
        "/api/references",
        files={"file": ("ref.txt", b"Extrait de style.", "text/plain")},
    )
    # embeddings de nouveau opérationnels : toujours rien à signaler
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.json()["style_indisponible"] == ""

    async def embed_ko(text, cfg):
        raise rag.EmbeddingUnavailable("Modèle d'embeddings « x » absent.")

    monkeypatch.setattr(rag, "embed", embed_ko)
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 200
    assert "absent" in r.json()["style_indisponible"]


def test_structure_verrouillage_pendant_analyse(client, monkeypatch, mock_embed):
    """Si le coffre se verrouille pendant l'analyse LLM, le résultat n'est
    plus jeté en 500 opaque : 423 explicite (l'UI ré-affiche l'écran de
    verrouillage et la dictée n'est pas perdue)."""
    async def structure_puis_verrou(*a, **k):
        security.lock()
        return {"updates": [], "questions": []}

    monkeypatch.setattr(llm, "structure", structure_puis_verrou)
    bid = client.post("/api/bilans", json={"domaines": []}).json()["id"]
    r = client.post(f"/api/bilans/{bid}/structure", json={"transcription": "Texte."})
    assert r.status_code == 423


# --- premier lancement guidé -------------------------------------------------------

def test_installation_etat(client, monkeypatch):
    from app import systeme

    monkeypatch.setattr(systeme, "ollama_etat", lambda cfg: {"ok": True, "modeles": ["x"]})
    etat = client.get("/api/installation").json()
    assert {"ollama", "ram_gio", "proposition", "pret"} <= set(etat)
    assert etat["ollama"] is True and etat["pret"] is False


def test_installation_accessible_verrouillee(client, monkeypatch):
    """L'écran d'installation doit fonctionner avant tout déverrouillage."""
    from app import systeme

    monkeypatch.setattr(systeme, "ollama_etat", lambda cfg: {"ok": False, "modeles": []})
    client.post("/api/lock")
    assert client.get("/api/installation").status_code == 200


def test_pull_nom_invalide(client):
    assert client.post(
        "/api/installation/pull", json={"modele": "méchant; rm -rf"}
    ).status_code == 400
    assert client.post("/api/installation/pull", json={}).status_code == 400


# --- dictée ----------------------------------------------------------------------------

def test_endpoints_legacy_supprimes(client):
    """Le trio legacy non verrouillé est retiré ; /api/models (utilisé par le
    sélecteur de l'interface) est conservé mais exige le déverrouillage."""
    assert client.post("/api/generate", json={"section": "anamnese", "notes": "x"}).status_code == 404
    assert client.get("/api/sections").status_code == 404


def test_models_exige_le_deverrouillage(client, monkeypatch):
    async def fake_models(host=None):
        return ["m1", "m2"]

    monkeypatch.setattr(llm, "list_models", fake_models)
    assert client.get("/api/models").json()["models"] == ["m1", "m2"]
    client.post("/api/lock")
    assert client.get("/api/models").status_code == 423


def test_models_suit_la_config_praticien(client, monkeypatch):
    """Le sélecteur interroge l'hôte Ollama *configuré* et annonce le modèle
    configuré : interroger la constante du module affichait sinon une liste
    sans rapport avec la configuration effective."""
    vus = {}

    async def fake_models(host=None):
        vus["host"] = host
        return ["m1"]

    monkeypatch.setattr(llm, "list_models", fake_models)
    client.put("/api/config", json={"overrides": {
        "llm": {"host": "http://127.0.0.1:11500", "model": "mon-modele"}
    }})
    assert client.get("/api/models").json()["default"] == "mon-modele"
    assert vus["host"] == "http://127.0.0.1:11500"


def test_transcribe_audio_vide(client):
    r = client.post("/api/transcribe", files={"audio": ("d.webm", b"", "audio/webm")})
    assert r.status_code == 400


def test_transcribe_echec_message_simple(client, monkeypatch):
    """La trace technique (ffmpeg, HuggingFace…) part dans le journal, jamais
    dans l'interface (UX-03)."""
    from app import stt

    def boom(data, filename, cfg):
        raise RuntimeError("ffmpeg error: /usr/lib/libavcodec.so introuvable")

    monkeypatch.setattr(stt, "transcribe", boom)
    r = client.post("/api/transcribe", files={"audio": ("d.webm", b"abc", "audio/webm")})
    assert r.status_code == 500
    detail = r.json()["detail"]
    assert "ffmpeg" not in detail and "libavcodec" not in detail
    assert "transcription a échoué" in detail


def test_stt_info(client):
    info = client.get("/api/stt/info").json()
    assert {"device", "compute_type", "model"} <= set(info)


# --- Bornes des corps envoyés (revue 2026-08-11, 5.5) -----------------------

def test_transcribe_audio_trop_volumineux_refuse(client, monkeypatch):
    """Le corps est lu par blocs et refusé (413) dès que le plafond est
    dépassé : rien n'est transmis au moteur de dictée."""
    from app import main as main_mod
    from app import stt

    monkeypatch.setattr(main_mod, "TAILLE_MAX_AUDIO", 2048)
    monkeypatch.setattr(main_mod, "_BLOC_LECTURE", 512)
    appels = []
    monkeypatch.setattr(stt, "transcribe", lambda *a: appels.append(a) or {"text": ""})
    gros = b"\0" * 4096
    r = client.post("/api/transcribe", files={"audio": ("d.webm", gros, "audio/webm")})
    assert r.status_code == 413
    assert "trop volumineux" in r.json()["detail"]
    assert appels == []
    # sous le plafond : accepté normalement
    r = client.post("/api/transcribe", files={"audio": ("d.webm", b"\0" * 1000, "audio/webm")})
    assert r.status_code == 200 and len(appels) == 1


def test_import_reference_trop_volumineux_refuse(client, monkeypatch, mock_embed):
    from app import importer
    from app import main as main_mod

    monkeypatch.setattr(main_mod, "TAILLE_MAX_DOCUMENT", 1024)
    appels = []
    monkeypatch.setattr(importer, "decouper", lambda *a: appels.append(a) or [])
    r = client.post(
        "/api/references",
        files={"file": ("bilan.txt", b"x" * 4096, "text/plain")},
        data={"domaine": ""},
    )
    assert r.status_code == 413
    assert appels == []
# --- Abandon d'une analyse par le navigateur (revue 2026-08-11, 2.4) --------

def test_analyse_abandonnee_quand_le_navigateur_ferme_la_requete(monkeypatch):
    """Le bouton « Annuler » ferme la requête : la tâche (donc l'appel au
    modèle) est annulée et rien ne sera écrit en base."""
    import asyncio

    import pytest

    from app import main as main_mod

    monkeypatch.setattr(main_mod, "_PAS_SURVEILLANCE_S", 0.01)
    etat = {"annulee": False, "sondages": 0}

    async def analyse_longue():
        try:
            await asyncio.sleep(30)
        except asyncio.CancelledError:
            etat["annulee"] = True
            raise
        return {"updates": []}

    class NavigateurQuiPart:
        async def is_disconnected(self):
            etat["sondages"] += 1
            return etat["sondages"] >= 2

    async def scenario():
        with pytest.raises(main_mod.AnalyseAbandonnee):
            await main_mod._jusqu_au_depart_du_client(NavigateurQuiPart(), analyse_longue())

    asyncio.run(scenario())
    assert etat["annulee"] is True and etat["sondages"] == 2


def test_analyse_terminee_rend_son_resultat(monkeypatch):
    import asyncio

    from app import main as main_mod

    monkeypatch.setattr(main_mod, "_PAS_SURVEILLANCE_S", 0.01)

    async def analyse_courte():
        await asyncio.sleep(0.03)
        return {"updates": [1]}

    class NavigateurPresent:
        async def is_disconnected(self):
            return False

    res = asyncio.run(main_mod._jusqu_au_depart_du_client(NavigateurPresent(), analyse_courte()))
    assert res == {"updates": [1]}


def test_structure_route_ne_sonde_pas_le_client_si_le_modele_repond(client, monkeypatch):
    """Chemin nominal : le résultat du modèle est persisté comme avant."""
    from app import llm

    async def structure_immediate(*a, **k):
        return {"updates": [{"section": "anamnese", "texte": "Plainte : lecture lente."}],
                "questions": [], "updates_non_placees": []}

    monkeypatch.setattr(llm, "structure", structure_immediate)
    b = client.post("/api/bilans", json={"domaines": []}).json()
    r = client.post(f"/api/bilans/{b['id']}/structure",
                    json={"transcription": "lecture lente", "reponses": []})
    assert r.status_code == 200
    sections = {s["cle"]: s for s in r.json()["bilan"]["sections"]}
    assert "lecture lente" in sections["anamnese"]["contenu"]
