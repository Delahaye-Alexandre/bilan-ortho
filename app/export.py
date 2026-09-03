"""Export d'un bilan : Markdown, texte, Word (.docx) et PDF.

Le compte-rendu de bilan orthophonique est un document adressé au médecin
prescripteur. Il portait jusqu'ici les seules rubriques : ni identité du
praticien, ni date, ni destinataire, ni signature — donc rien qui permette de
l'envoyer tel quel. Le praticien recollait le texte dans son papier à en-tête,
ce qui reprenait le temps que l'outil venait de faire gagner.

L'en-tête, le destinataire et la signature n'apparaissent que s'ils ont été
renseignés : aucune identité n'est inventée, et un coffre neuf produit le même
document qu'avant.

La mise en page (police, corps, interligne, marges, couleur des titres,
numérotation des rubriques, numéros de page, logo) vient de la section
`mise_en_page` de la configuration (lot B du plan « mise en forme ») : le
Word et le PDF la lisent tous deux, l'écran Paramètres en montre l'effet sur
un bilan fictif (`bilan_exemple`).
"""
from __future__ import annotations

import base64
import io
import os
from datetime import date
from pathlib import Path

from . import config, cotation, texte_riche
from .bilan import DRAPEAU_LIBELLE, etalonnage_texte
from .patient import age_texte, date_fr

DISCLAIMER = (
    "Document généré comme aide à la rédaction, relu et validé par l'orthophoniste : "
    "la responsabilité du contenu lui revient entièrement. "
    "L'outil ne pose aucun diagnostic."
)

# Tant que le bilan n'est pas marqué validé, le document le dit lui-même : sans
# cette mention, l'export d'un brouillon non relu était indiscernable du
# compte-rendu définitif — exactement le geste qu'un outil médico-légal doit
# rendre difficile.
MENTION_BROUILLON = (
    "BROUILLON — document de travail non validé, à ne pas transmettre en l'état."
)

_TYPE_LBL = {
    "initial_simple": "Bilan initial (simple)",
    "initial_complexe": "Bilan initial (complexe)",
    "renouvellement": "Bilan de renouvellement",
}

_COLONNES = ["Test", "Épreuve", "Score brut", "Étalonnage", "Interprétation"]

# Titres déjà porteurs de civilité : n'en ajoutons pas une seconde.
_CIVILITES = ("dr", "dr.", "docteur", "pr", "pr.", "professeur", "mme", "m.", "mr")


def _naissance(p: dict) -> str:
    """Mention de naissance accordée au sexe *enregistré*, sans parenthèse.

    Le sexe est une donnée du dossier : quand il est connu, l'accorder est plus
    juste qu'un « né(e) » parenthésé dans un document adressé au prescripteur.
    Quand il ne l'est pas (non renseigné, ou « autre »), la date est introduite
    sans participe plutôt que par une forme genrée par défaut."""
    date = date_fr(p["date_naissance"])
    participe = {"F": "née", "M": "né"}.get((p.get("sexe") or "").strip().upper())
    return f"{participe} le {date}" if participe else f"date de naissance : {date}"


def _date_reference(b: dict) -> str:
    """Date du bilan, à défaut sa date de création.

    L'âge doit être calculé à la date du bilan : un compte-rendu rédigé une
    semaine plus tard ne doit pas vieillir le patient d'autant."""
    return (b.get("date_bilan") or "").strip() or (b.get("created_at") or "")


def _patient_ligne(b: dict) -> str:
    p = b.get("patient")
    if not p:
        return ""
    ident = " ".join(x for x in [(p.get("nom") or "").upper(), p.get("prenom") or ""] if x)
    ligne = f"Patient : {ident or '—'}"
    if p.get("date_naissance"):
        ligne += f", {_naissance(p)}"
        age = age_texte(p["date_naissance"], _date_reference(b))
        if age:
            ligne += f" ({age} à la date du bilan)"
    return ligne


def _nom_praticien(prat: dict) -> str:
    return " ".join(
        x for x in [(prat.get("prenom") or "").strip(), (prat.get("nom") or "").strip()] if x
    )


def _entete_praticien(prat: dict) -> list[str]:
    """Lignes d'en-tête du cabinet, dans l'ordre d'un papier à lettres.

    Sans nom, pas d'en-tête du tout : `titre` vaut « Orthophoniste » par défaut,
    et le seul métier ne constitue pas une identité — un coffre neuf produirait
    un en-tête qui n'identifie personne."""
    nom = _nom_praticien(prat)
    if not nom:
        return []
    lignes = [nom]
    if (prat.get("titre") or "").strip():
        lignes.append(prat["titre"].strip())
    ville = " ".join(
        x for x in [(prat.get("code_postal") or "").strip(), (prat.get("ville") or "").strip()] if x
    )
    adresse = ", ".join(x for x in [(prat.get("adresse") or "").strip(), ville] if x)
    if adresse:
        lignes.append(adresse)
    contact = " — ".join(
        x for x in [
            f"Tél. {prat['telephone'].strip()}" if (prat.get("telephone") or "").strip() else "",
            (prat.get("email") or "").strip(),
        ] if x
    )
    if contact:
        lignes.append(contact)
    ids = " — ".join(
        x for x in [
            f"N° ADELI {prat['adeli'].strip()}" if (prat.get("adeli") or "").strip() else "",
            f"RPPS {prat['rpps'].strip()}" if (prat.get("rpps") or "").strip() else "",
            f"SIRET {prat['siret'].strip()}" if (prat.get("siret") or "").strip() else "",
        ] if x
    )
    if ids:
        lignes.append(ids)
    return lignes


def _destinataire(b: dict) -> str:
    nom = ((b.get("prescripteur") or {}).get("nom") or "").strip()
    if not nom:
        return ""
    premier = nom.split()[0].lower()
    if premier in _CIVILITES:
        return f"À l'attention de {nom}"
    return f"À l'attention du Dr {nom}"


def _signature(b: dict, prat: dict) -> list[str]:
    """« Fait à …, le … » puis l'identité qui signe. Rien sans identité."""
    nom = _nom_praticien(prat)
    if not nom:
        return []
    lieu = (prat.get("lieu_signature") or "").strip() or (prat.get("ville") or "").strip()
    quand = date_fr(_date_reference(b)[:10]) if _date_reference(b) else ""
    formule = ", le ".join(x for x in [f"Fait à {lieu}" if lieu else "Fait", quand] if x)
    titre = (prat.get("titre") or "").strip()
    return [formule, f"{nom}, {titre.lower()}" if titre else nom]


def _table_epreuves(b: dict) -> list[list[str]] | None:
    """Lignes du tableau des résultats, ou None s'il n'y a rien à montrer."""
    lignes: list[list[str]] = []
    for e in b.get("epreuves") or []:
        for r in e.get("resultats") or []:
            interpretation = " ".join(
                x for x in [
                    DRAPEAU_LIBELLE.get(r.get("drapeau_seuil") or "", ""),
                    (r.get("interpretation") or "").strip(),
                ] if x
            )
            lignes.append([
                e.get("test_nom") or "",
                r.get("sous_epreuve") or "",
                r.get("score_brut") or "",
                etalonnage_texte(r),
                interpretation,
            ])
    return lignes or None


def mise_en_page(cfg: dict | None) -> dict:
    """Réglages de mise en page effectifs : les défauts, complétés par ce que
    porte la configuration (qui peut être partielle ou absente : les tests et
    les anciens appels passent `cfg=None`)."""
    return config._deep_merge(
        config.DEFAULTS["mise_en_page"], (cfg or {}).get("mise_en_page") or {}
    )


def _logo_octets(mp: dict) -> bytes | None:
    """Octets de l'image du logo, ou None (absent ou illisible : le document
    sort sans logo plutôt que de ne pas sortir)."""
    logo = mp.get("logo")
    if not isinstance(logo, dict) or not logo.get("donnees"):
        return None
    try:
        return base64.b64decode(logo["donnees"], validate=True)
    except (ValueError, TypeError):
        return None


# Un logo est une petite image d'en-tête : 400 px de haut suffisent largement
# pour 20 mm imprimés (≈ 500 ppp). Au-delà, l'image est réduite avant d'être
# rangée en base64 dans la configuration chiffrée, jamais en fichier à côté.
HAUTEUR_LOGO_PX = 400
TAILLE_MAX_LOGO = 3 * 1024 * 1024


def preparer_logo(data: bytes) -> dict:
    """Vérifie (Pillow, pas l'extension), réduit et ré-encode un logo déposé.

    Retourne l'entrée `mise_en_page.logo` de la configuration. ValueError si
    ce n'est pas un PNG ou un JPEG lisible."""
    from PIL import Image

    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception as exc:
        raise ValueError("Image illisible : déposez un fichier PNG ou JPEG.") from exc
    fmt = img.format
    if fmt not in ("PNG", "JPEG"):
        raise ValueError(
            f"Format {fmt or 'inconnu'} non pris en charge : déposez un PNG ou un JPEG."
        )
    if fmt == "JPEG":
        img = img.convert("RGB")
    elif img.mode not in ("RGB", "RGBA"):
        # Palette ou niveaux de gris : la réduction d'une image en palette se
        # fait sans lissage ; en RGBA la transparence est conservée.
        img = img.convert("RGBA")
    if img.height > HAUTEUR_LOGO_PX:
        largeur = max(1, round(img.width * HAUTEUR_LOGO_PX / img.height))
        img = img.resize((largeur, HAUTEUR_LOGO_PX), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    if fmt == "JPEG":
        img.save(buf, "JPEG", quality=88)
        mime = "image/jpeg"
    else:
        img.save(buf, "PNG", optimize=True)
        mime = "image/png"
    return {
        "type": mime,
        "donnees": base64.b64encode(buf.getvalue()).decode("ascii"),
        "largeur": img.width,
        "hauteur": img.height,
    }


def bilan_exemple(cfg: dict | None = None) -> dict:
    """Bilan fictif, court, qui passe par tous les blocs du document (en-tête,
    destinataire, rubriques en texte riche, tableau, cotation, signature) :
    l'aperçu de l'écran Paramètres le met en page avec les réglages en cours.
    Aucune donnée réelle ; la cotation suit la configuration."""
    cot = ((cfg or {}).get("cotation")) or config.DEFAULTS["cotation"]
    coeff = float(cot["bilan_simple_coeff"])
    valeur = float(cot["valeur_amo"])
    aujourd_hui = date.today().isoformat()
    return {
        "id": 0,
        "type": "initial_simple",
        "statut": "valide",
        "domaine_titres": "Langage oral",
        "date_bilan": aujourd_hui,
        "created_at": aujourd_hui,
        "patient": {"nom": "Exemple", "prenom": "Camille",
                    "date_naissance": "2018-03-12", "sexe": ""},
        "prescripteur": {"nom": "Dr Exemple"},
        "sections": [
            {"titre": "Anamnèse", "contenu": (
                "Bilan demandé par le médecin traitant pour des difficultés de "
                "langage signalées à l'école. Pas d'antécédent ORL rapporté ; "
                "audition contrôlée en **mars 2026**."
            )},
            {"titre": "Observations", "contenu": (
                "Comportement attentif et coopérant pendant la passation.\n"
                "- <u>Compréhension orale</u> : consignes simples et complexes suivies.\n"
                "- <u>Expression</u> : phrases courtes, lexique réduit, *quelques "
                "simplifications phonologiques*."
            )},
            {"titre": "Conclusion et projet", "contenu": (
                "Les observations recueillies justifient un accompagnement "
                "orthophonique.\n1. Enrichissement du lexique.\n2. Travail "
                "phonologique.\n3. Réévaluation dans six mois."
            )},
        ],
        "epreuves": [{"test_nom": "Test fictif", "resultats": [
            {"sous_epreuve": "Dénomination", "score_brut": "18/30",
             "etalonnage_type": "ecart_type", "etalonnage_valeur": "-1,2",
             "drapeau_seuil": "fragilite", "interpretation": ""},
            {"sous_epreuve": "Compréhension", "score_brut": "27/30",
             "etalonnage_type": "ecart_type", "etalonnage_valeur": "0,3",
             "drapeau_seuil": "norme", "interpretation": ""},
        ]}],
        "cotation": {"coefficient": coeff, "montant": round(coeff * valeur, 2),
                     "valeur_lettre_cle": valeur},
    }


def est_brouillon(b: dict) -> bool:
    """Un bilan qui n'a pas été explicitement validé (ou envoyé) reste un
    brouillon — y compris un bilan vide, qui s'exportait jusqu'ici avec
    en-tête, date et bloc de signature."""
    return (b.get("statut") or "brouillon") not in ("valide", "envoye")


def _content(b: dict, cfg: dict | None = None) -> list[tuple[str, object]]:
    """Document sous forme de blocs (type, contenu), rendus par chaque format."""
    prat = ((cfg or {}).get("praticien")) or {}
    mp = mise_en_page(cfg)
    blocks: list[tuple[str, object]] = []
    if est_brouillon(b):
        blocks.append(("brouillon", MENTION_BROUILLON))
    logo = _logo_octets(mp)
    if logo:
        # Word et PDF seulement ; Markdown et texte l'ignorent.
        blocks.append(("logo", {
            "donnees": logo, "position": mp["logo_position"],
            "hauteur_mm": float(mp["logo_hauteur_mm"]),
        }))
    entete = _entete_praticien(prat)
    if entete:
        blocks.append(("entete", entete))
    dest = _destinataire(b)
    if dest:
        blocks.append(("dest", dest))
    doms = b.get("domaine_titres") or "Générique"
    blocks.append(("h1", "Compte-rendu de bilan orthophonique"))
    blocks.append(("p", f"{_TYPE_LBL.get(b.get('type'), b.get('type', ''))} · Domaine(s) : {doms}"))
    patient = _patient_ligne(b)
    if patient:
        blocks.append(("p", patient))
    if (b.get("date_bilan") or "").strip():
        blocks.append(("p", f"Date du bilan : {date_fr(b['date_bilan'])}"))
    for s in b.get("sections", []):
        if (s.get("contenu") or "").strip():
            blocks.append(("h2", s["titre"]))
            # Contenu de rubrique : texte riche (gras, listes…) — voir texte_riche.
            blocks.append(("riche", s["contenu"].strip()))
    table = _table_epreuves(b)
    if table:
        blocks.append(("h2", "Résultats des épreuves"))
        blocks.append(("table", table))
    cot = b.get("cotation")
    if cot:
        # Le code est reconstruit depuis le coefficient plutôt que repris tel
        # quel : les cotations déjà enregistrées portent « AMO 24.0 ». C'est le
        # seul chiffre du document que l'Assurance maladie peut recouper.
        coeff = cotation.coeff_texte(cot["coefficient"])
        blocks.append(("h2", "Cotation (NGAP)"))
        blocks.append((
            "p",
            f"AMO {coeff} — {cotation.euros(cot['montant'])} "
            f"(coefficient {coeff}, valeur lettre-clé "
            f"{cotation.euros(cot['valeur_lettre_cle'])})",
        ))
    sign = _signature(b, prat)
    if sign:
        blocks.append(("sign", sign))
    blocks.append(("hr", ""))
    blocks.append(("i", DISCLAIMER))
    if mp.get("rubriques_numerotees"):
        # Rubriques, résultats et cotation numérotés à la suite : une
        # numérotation qui s'arrêterait aux seules rubriques se lirait comme
        # un oubli sur le tableau des résultats.
        n = 0
        for i, (k, t) in enumerate(blocks):
            if k == "h2":
                n += 1
                blocks[i] = ("h2", f"{n}. {t}")
    return blocks


def _cellule_md(v: str) -> str:
    """Cellule de tableau Markdown : ni « | » ni retour à la ligne.

    Une interprétation dictée en contenant cassait la table entière — et le
    .md est le format par défaut de la route d'export."""
    v = (v or "").replace("|", "\\|").replace("\r", " ").replace("\n", " ").strip()
    return v or "—"


def _cellule_txt(v: str) -> str:
    """Cellule de tableau en texte brut : les colonnes sont alignées au
    caractère, un retour à la ligne décalerait tout ce qui suit."""
    v = (v or "").replace("\r", " ").replace("\n", " ").strip()
    return v or "—"


def to_markdown(b: dict, cfg: dict | None = None) -> str:
    out: list[str] = []
    for k, t in _content(b, cfg):
        if k == "brouillon":
            out.append(f"> **{t}**")
        elif k == "entete":
            out.append("\n".join(f"**{ligne}**" if i == 0 else ligne
                                 for i, ligne in enumerate(t)))
            out.append("\n---")
        elif k == "dest":
            out.append(f"\n{t}")
        elif k == "h1":
            out.append(f"\n# {t}")
        elif k == "h2":
            out.append(f"\n## {t}")
        elif k == "p":
            out.append(f"\n{t}")
        elif k == "riche":
            out.append("\n" + texte_riche.canonique(t))
        elif k == "table":
            out.append("\n| " + " | ".join(_COLONNES) + " |")
            out.append("| " + " | ".join("---" for _ in _COLONNES) + " |")
            for ligne in t:
                out.append("| " + " | ".join(_cellule_md(c) for c in ligne) + " |")
        elif k == "sign":
            out.append("\n" + "  \n".join(t))
        elif k == "hr":
            out.append("\n---")
        elif k == "i":
            out.append(f"\n_{t}_")
    return "\n".join(out).strip() + "\n"


def to_txt(b: dict, cfg: dict | None = None) -> str:
    out: list[str] = []
    for k, t in _content(b, cfg):
        if k == "brouillon":
            out.append(t)
            out.append("=" * 40)
        elif k == "entete":
            out.extend(t)
            out.append("-" * 40)
        elif k == "dest":
            out.append("\n" + t)
        elif k in ("h1", "h2"):
            out.append("\n" + t.upper())
        elif k == "p":
            out.append(t)
        elif k == "riche":
            out.append(texte_riche.en_clair(t))
        elif k == "table":
            cellules = [[_cellule_txt(c) for c in ligne] for ligne in t]
            larg = [
                max(len(_COLONNES[i]), *(len(ligne[i]) for ligne in cellules))
                for i in range(len(_COLONNES))
            ]
            out.append("  ".join(c.ljust(larg[i]) for i, c in enumerate(_COLONNES)))
            out.append("  ".join("-" * w for w in larg))
            for ligne in cellules:
                out.append("  ".join(c.ljust(larg[i]) for i, c in enumerate(ligne)))
        elif k == "sign":
            out.append("")
            out.extend(t)
        elif k == "hr":
            out.append("-" * 40)
        elif k == "i":
            out.append(t)
    return "\n".join(out).strip() + "\n"


def _docx_runs(paragraphe, segments: list[texte_riche.Segment]) -> None:
    """Segments -> runs Word (gras, italique, souligné, retours à la ligne)."""
    for s in segments:
        for i, morceau in enumerate(s.texte.split("\n")):
            if i:
                paragraphe.add_run().add_break()
            if not morceau:
                continue
            run = paragraphe.add_run(morceau)
            if s.gras:
                run.bold = True
            if s.italique:
                run.italic = True
            if s.souligne:
                run.underline = True


def _docx_numerotation_neuve(doc, nom_style: str):
    """Nouvelle instance de numérotation repartant à 1, calquée sur le style.

    python-docx fait partager à tous les paragraphes « List Number » une même
    numérotation : la deuxième liste du document continuait à 4. Retourne
    l'identifiant à poser sur les paragraphes, ou None si le document (un
    gabarit du praticien, par exemple) ne s'y prête pas — la liste continue
    alors, plutôt que d'échouer."""
    try:
        from docx.oxml.ns import qn

        style = doc.styles[nom_style]
        num_id = style.element.pPr.numPr.numId.val
        numbering = doc.part.numbering_part.element
        origine = numbering.num_having_numId(num_id)
        neuf = numbering.add_num(origine.abstractNumId.val)
        neuf.add_lvlOverride(ilvl=0).add_startOverride(1)
        qn("w:numId")  # import utilisé : garde le linter tranquille
        return neuf.numId
    except Exception:
        return None


def _docx_liste(doc, bloc: texte_riche.Liste) -> None:
    nom_style = "List Number" if bloc.ordonnee else "List Bullet"
    num_id = _docx_numerotation_neuve(doc, nom_style) if bloc.ordonnee else None
    for i, item in enumerate(bloc.items, 1):
        try:
            p = doc.add_paragraph(style=nom_style)
        except KeyError:
            # Style absent (gabarit personnalisé) : le marqueur devient du texte.
            p = doc.add_paragraph(f"{i}. " if bloc.ordonnee else "- ")
        else:
            if num_id is not None:
                num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
                num_pr.get_or_add_numId().val = num_id
                num_pr.get_or_add_ilvl().val = 0
        _docx_runs(p, item)


def _docx_riche(doc, texte: str) -> None:
    for bloc in texte_riche.analyser(texte):
        if isinstance(bloc, texte_riche.Liste):
            _docx_liste(doc, bloc)
        else:
            _docx_runs(doc.add_paragraph(), bloc.segments)


_ALIGNEMENTS = {"gauche": "LEFT", "centre": "CENTER", "droite": "RIGHT"}


def _docx_police(style, nom: str) -> None:
    """Pose une police sur un style, en retirant les renvois au thème : sinon
    Word garde « Calibri Light » (titres du thème) quel que soit le nom posé."""
    from docx.oxml.ns import qn

    style.font.name = nom
    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        rfonts.attrib.pop(qn(attr), None)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), nom)


def _docx_champ(paragraphe, instruction: str) -> None:
    """Champ Word simple (PAGE, NUMPAGES…) : calculé par Word à l'ouverture."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    champ = OxmlElement("w:fldSimple")
    champ.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    texte = OxmlElement("w:t")
    texte.text = "1"
    run.append(texte)
    champ.append(run)
    paragraphe._p.append(champ)


def _docx_mise_en_page(doc, mp: dict) -> None:
    """Applique police, corps, interligne, marges, couleur des titres et
    numéros de page au document (styles Normal, Heading 1 et Heading 2, et
    sections). Le lot D (gabarit .docx du praticien) court-circuitera cette
    fonction : les styles viendront alors du gabarit."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, RGBColor

    taille = float(mp["taille_corps"])
    normal = doc.styles["Normal"]
    _docx_police(normal, mp["police"])
    normal.font.size = Pt(taille)
    normal.paragraph_format.line_spacing = float(mp["interligne"])
    couleur = RGBColor.from_string(mp["couleur_titres"].lstrip("#").upper())
    for nom, delta in (("Heading 1", 6), ("Heading 2", 2)):
        st = doc.styles[nom]
        _docx_police(st, mp["police"])
        st.font.size = Pt(taille + delta)
        st.font.bold = True
        st.font.color.rgb = couleur
    marge = Mm(float(mp["marges_mm"]))
    for section in doc.sections:
        section.left_margin = section.right_margin = marge
        section.top_margin = section.bottom_margin = marge
        if mp.get("numeros_de_page"):
            pied = section.footer
            p = pied.paragraphs[0] if pied.paragraphs else pied.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("Page ")
            _docx_champ(p, "PAGE")
            p.add_run(" / ")
            _docx_champ(p, "NUMPAGES")


def to_docx(b: dict, cfg: dict | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    doc = Document()
    _docx_mise_en_page(doc, mise_en_page(cfg))
    for k, t in _content(b, cfg):
        if k == "brouillon":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(t).bold = True
        elif k == "logo":
            try:
                doc.add_picture(io.BytesIO(t["donnees"]), height=Mm(t["hauteur_mm"]))
            except Exception:
                continue  # image inattendue : le document sort sans logo
            doc.paragraphs[-1].alignment = getattr(WD_ALIGN_PARAGRAPH, _ALIGNEMENTS[t["position"]])
        elif k == "entete":
            for i, ligne in enumerate(t):
                p = doc.add_paragraph()
                run = p.add_run(ligne)
                run.bold = i == 0
                run.font.size = None if i == 0 else run.font.size
        elif k == "dest":
            p = doc.add_paragraph(t)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif k == "h1":
            doc.add_heading(t, level=1)
        elif k == "h2":
            doc.add_heading(t, level=2)
        elif k == "p":
            doc.add_paragraph(t)
        elif k == "riche":
            _docx_riche(doc, t)
        elif k == "table":
            table = doc.add_table(rows=1, cols=len(_COLONNES))
            table.style = "Table Grid"
            for i, titre in enumerate(_COLONNES):
                table.rows[0].cells[i].paragraphs[0].add_run(titre).bold = True
            for ligne in t:
                cells = table.add_row().cells
                for i, val in enumerate(ligne):
                    cells[i].text = val or "—"
        elif k == "sign":
            doc.add_paragraph()
            for ligne in t:
                p = doc.add_paragraph(ligne)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Espace laissé à la signature manuscrite ou au cachet.
            doc.add_paragraph()
            doc.add_paragraph()
        elif k == "hr":
            doc.add_paragraph("_" * 30)
        elif k == "i":
            p = doc.add_paragraph()
            p.add_run(t).italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Fichiers TrueType des polices proposées dans Paramètres (régulier, gras,
# italique, gras italique), tels que Windows les nomme. Trouvés sur la
# machine, ils sont incorporés au PDF ; sinon le PDF prend la police intégrée
# équivalente (Helvetica sans empattements, Times avec).
_FICHIERS_POLICES = {
    "arial": ("arial.ttf", "arialbd.ttf", "ariali.ttf", "arialbi.ttf"),
    "calibri": ("calibri.ttf", "calibrib.ttf", "calibrii.ttf", "calibriz.ttf"),
    "verdana": ("verdana.ttf", "verdanab.ttf", "verdanai.ttf", "verdanaz.ttf"),
    "times new roman": ("times.ttf", "timesbd.ttf", "timesi.ttf", "timesbi.ttf"),
    "georgia": ("georgia.ttf", "georgiab.ttf", "georgiai.ttf", "georgiaz.ttf"),
}
_SERIF = ("times", "georgia", "cambria", "garamond", "book antiqua", "palatino", "serif")
_FACES = ("normal", "gras", "italique", "gras_italique")
_POLICES_PDF: dict[str, dict[str, str]] = {}
_CHEMINS_POLICES: dict[str, Path | None] = {}


def _dossiers_polices() -> list[Path]:
    dossiers: list[Path] = []
    windir = os.environ.get("WINDIR") or os.environ.get("SystemRoot")
    if windir:
        dossiers.append(Path(windir) / "Fonts")
    maison = Path.home()
    dossiers += [
        maison / "AppData" / "Local" / "Microsoft" / "Windows" / "Fonts",
        Path("/usr/share/fonts"), Path("/usr/local/share/fonts"),
        maison / ".fonts", maison / ".local" / "share" / "fonts",
        Path("/Library/Fonts"), Path("/System/Library/Fonts"),
    ]
    return [d for d in dossiers if d.is_dir()]


def _trouver_police(nom_fichier: str) -> Path | None:
    if nom_fichier in _CHEMINS_POLICES:
        return _CHEMINS_POLICES[nom_fichier]
    trouve = None
    for dossier in _dossiers_polices():
        direct = dossier / nom_fichier
        if direct.is_file():
            trouve = direct
            break
        try:
            trouve = next(
                (f for f in dossier.rglob("*") if f.name.lower() == nom_fichier and f.is_file()),
                None,
            )
        except OSError:
            trouve = None
        if trouve:
            break
    _CHEMINS_POLICES[nom_fichier] = trouve
    return trouve


def _polices_pdf(famille: str) -> dict[str, str]:
    """Noms de police reportlab (normal, gras, italique, gras italique) pour
    une famille choisie dans Paramètres ; la famille TrueType est enregistrée
    au premier usage pour que <b> et <i> du balisage y renvoient."""
    cle = (famille or "").strip().lower()
    if cle in _POLICES_PDF:
        return _POLICES_PDF[cle]
    if any(s in cle for s in _SERIF):
        repli = dict(zip(_FACES, ("Times-Roman", "Times-Bold", "Times-Italic", "Times-BoldItalic")))
    else:
        repli = dict(zip(_FACES, ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique",
                                  "Helvetica-BoldOblique")))
    resultat = repli
    fichiers = _FICHIERS_POLICES.get(cle)
    chemins = [_trouver_police(f) for f in fichiers] if fichiers else []
    if chemins and chemins[0] is not None:
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            prefixe = "MEP-" + cle.replace(" ", "")
            noms = []
            for suffixe, chemin in zip(("", "-Bold", "-Italic", "-BoldItalic"), chemins):
                nom = prefixe + suffixe
                # Face absente (gras seul manquant…) : la régulière la remplace.
                pdfmetrics.registerFont(TTFont(nom, str(chemin or chemins[0])))
                noms.append(nom)
            pdfmetrics.registerFontFamily(
                prefixe, normal=noms[0], bold=noms[1], italic=noms[2], boldItalic=noms[3],
            )
            resultat = dict(zip(_FACES, noms))
        except Exception:
            resultat = repli
    _POLICES_PDF[cle] = resultat
    return resultat


def _canvas_numerote(police: str, taille: float, y: float):
    """Fabrique de canvas reportlab qui écrit « Page i / n » en pied de page.
    Le total n n'est connu qu'à la fin : les pages sont mises en réserve et
    dessinées dans save()."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    class CanvasNumerote(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._pages_en_reserve: list[dict] = []

        def showPage(self):
            self._pages_en_reserve.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._pages_en_reserve)
            for i, etat in enumerate(self._pages_en_reserve, 1):
                self.__dict__.update(etat)
                self.saveState()
                self.setFont(police, taille)
                self.setFillGray(0.4)
                self.drawCentredString(A4[0] / 2, y, f"Page {i} / {total}")
                self.restoreState()
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

    return CanvasNumerote


def to_pdf(b: dict, cfg: dict | None = None) -> bytes:
    """PDF paginé, format d'envoi habituel d'un compte-rendu au prescripteur.

    reportlab est retenu pour rester compatible avec l'exécutable Windows :
    c'est une bibliothèque Python pure, sans moteur de rendu HTML ni binaire
    système à embarquer."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        Image,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    mp = mise_en_page(cfg)
    polices = _polices_pdf(mp["police"])
    taille = float(mp["taille_corps"])
    interligne = round(taille * 1.2 * float(mp["interligne"]), 1)
    marge = float(mp["marges_mm"]) * mm
    largeur_utile = A4[0] - 2 * marge
    couleur_titres = colors.HexColor(mp["couleur_titres"])

    ss = getSampleStyleSheet()
    corps = ParagraphStyle(
        "corps", parent=ss["BodyText"], fontName=polices["normal"],
        fontSize=taille, leading=interligne,
    )
    petit = ParagraphStyle(
        "petit", parent=corps, fontSize=taille - 1.5, leading=round((taille - 1.5) * 1.3, 1),
    )
    droite = ParagraphStyle("droite", parent=corps, alignment=TA_RIGHT)
    centre = ParagraphStyle("centre", parent=corps, alignment=TA_CENTER)
    titre1 = ParagraphStyle(
        "t1", parent=ss["Heading1"], fontName=polices["gras"], fontSize=taille + 6,
        leading=round((taille + 6) * 1.2, 1), spaceBefore=10, textColor=couleur_titres,
    )
    titre2 = ParagraphStyle(
        "t2", parent=ss["Heading2"], fontName=polices["gras"], fontSize=taille + 2,
        leading=round((taille + 2) * 1.2, 1), spaceBefore=10, textColor=couleur_titres,
    )

    def esc(txt: str) -> str:
        return (txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def balise(segments: list[texte_riche.Segment]) -> str:
        """Segments -> balisage de paragraphe reportlab (texte échappé d'abord :
        seules NOS balises passent)."""
        parts = []
        for s in segments:
            t = esc(s.texte).replace("\n", "<br/>")
            if s.souligne:
                t = f"<u>{t}</u>"
            if s.italique:
                t = f"<i>{t}</i>"
            if s.gras:
                t = f"<b>{t}</b>"
            parts.append(t)
        return "".join(parts)

    def riche(texte: str) -> list:
        flow: list = []
        for bloc in texte_riche.analyser(texte):
            if isinstance(bloc, texte_riche.Paragraphe):
                flow.append(Paragraph(balise(bloc.segments), corps))
                continue
            items = [ListItem(Paragraph(balise(i), corps), leftIndent=14) for i in bloc.items]
            flow.append(ListFlowable(
                items, bulletType="1" if bloc.ordonnee else "bullet",
                start=1 if bloc.ordonnee else "•", leftIndent=14,
                bulletFontSize=9, spaceBefore=2, spaceAfter=2,
            ))
        return flow

    def construire(tableau_en_lignes: bool) -> list:
        """Blocs du document. `tableau_en_lignes` rend les résultats sous forme
        de paragraphes plutôt que de tableau : c'est le repli quand une cellule
        est plus haute qu'une page (reportlab ne sait pas la découper)."""
        flow: list = []
        for k, t in _content(b, cfg):
            if k == "brouillon":
                flow.append(Paragraph(f"<b>{esc(t)}</b>", centre))
                flow.append(Spacer(1, 6))
            elif k == "logo":
                try:
                    img = Image(io.BytesIO(t["donnees"]))
                    hauteur = t["hauteur_mm"] * mm
                    largeur = hauteur * img.imageWidth / img.imageHeight
                    if largeur > largeur_utile:
                        hauteur, largeur = hauteur * largeur_utile / largeur, largeur_utile
                    img.drawWidth, img.drawHeight = largeur, hauteur
                    img.hAlign = _ALIGNEMENTS[t["position"]]
                except Exception:
                    continue  # image inattendue : le document sort sans logo
                flow.append(img)
                flow.append(Spacer(1, 6))
            elif k == "entete":
                for i, ligne in enumerate(t):
                    flow.append(Paragraph(
                        f"<b>{esc(ligne)}</b>" if i == 0 else esc(ligne),
                        corps if i == 0 else petit,
                    ))
                flow.append(Spacer(1, 4))
                flow.append(HRFlowable(width="100%", color=colors.grey))
            elif k == "dest":
                flow.append(Spacer(1, 8))
                flow.append(Paragraph(esc(t), droite))
            elif k == "h1":
                flow.append(Paragraph(esc(t), titre1))
            elif k == "h2":
                flow.append(Paragraph(esc(t), titre2))
            elif k == "p":
                flow.append(Paragraph(esc(t).replace("\n", "<br/>"), corps))
            elif k == "riche":
                # Les rubriques portent gras, listes et sauts de ligne signifiants.
                flow.extend(riche(t))
            elif k == "table" and tableau_en_lignes:
                for ligne in t:
                    parts = [f"<b>{esc(ligne[0])}</b>"] + [
                        f"{esc(_COLONNES[i])} : {esc(v)}"
                        for i, v in enumerate(ligne) if i and v
                    ]
                    flow.append(Paragraph(" — ".join(parts), petit))
            elif k == "table":
                data = [[Paragraph(f"<b>{esc(c)}</b>", petit) for c in _COLONNES]]
                data += [[Paragraph(esc(c or "—"), petit) for c in ligne] for ligne in t]
                # Colonnes au prorata de la largeur utile : des largeurs fixes
                # débordaient de la page dès que les marges s'élargissaient.
                parts = (32, 38, 28, 28, 40)
                table = Table(
                    data, colWidths=[largeur_utile * p / sum(parts) for p in parts],
                    repeatRows=1,
                )
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                flow.append(table)
            elif k == "sign":
                flow.append(Spacer(1, 16))
                for ligne in t:
                    flow.append(Paragraph(esc(ligne), droite))
                flow.append(Spacer(1, 24))  # place pour la signature ou le cachet
            elif k == "hr":
                flow.append(Spacer(1, 8))
                flow.append(HRFlowable(width="100%", color=colors.grey))
            elif k == "i":
                flow.append(Paragraph(f"<i>{esc(t)}</i>", petit))
        return flow

    def filigrane(canvas, doc):
        """« BROUILLON » en travers de chaque page tant que le bilan n'est pas
        validé : la mention doit rester lisible même sur une page imprimée
        isolément du reste du document."""
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 64)
        canvas.setFillGray(0.85)
        canvas.translate(A4[0] / 2, A4[1] / 2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, "BROUILLON")
        canvas.restoreState()

    def document(buf):
        return SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=marge, rightMargin=marge, topMargin=marge, bottomMargin=marge,
            title="Compte-rendu de bilan orthophonique",
        )

    pages = {"onFirstPage": filigrane, "onLaterPages": filigrane} if est_brouillon(b) else {}
    if mp.get("numeros_de_page"):
        pages["canvasmaker"] = _canvas_numerote(polices["normal"], taille - 2.5, marge / 2)
    buf = io.BytesIO()
    try:
        document(buf).build(construire(False), **pages)
    except Exception:
        # Une cellule plus haute qu'une page fait échouer toute la mise en page
        # (LayoutError) : une interprétation clinique un peu longue suffisait à
        # rendre le PDF impossible. Le document part quand même, résultats
        # rendus en lignes — mieux vaut une mise en forme dégradée qu'un export
        # refusé au moment de l'envoyer au prescripteur.
        buf = io.BytesIO()
        document(buf).build(construire(True), **pages)
    return buf.getvalue()
